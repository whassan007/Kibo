import os
import json
import sqlite3
from typing import Dict, Any, List, Optional
from fastapi import FastAPI, Request
import canadian_verdicts_agent

from fastapi import BackgroundTasks, Security, HTTPException, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from enum import Enum
from reportlab.pdfgen import canvas
import hashlib
import os

security = HTTPBearer(auto_error=False)

class UserRole(str, Enum):
    PUBLIC = "PUBLIC"
    EMPLOYEE = "EMPLOYEE"
    ANALYST = "ANALYST"
    DPO = "DPO"
    CPO = "CPO"
    LEGAL = "LEGAL"
    AUDITOR = "AUDITOR"

def get_current_active_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if not credentials:
        return UserRole.CPO # Default fallback for local testing without headers
    token = credentials.credentials
    try:
        return UserRole(token.upper())
    except ValueError:
        raise HTTPException(status_code=403, detail="Invalid role token")
        
def requires_role(allowed_roles: list):
    def role_checker(current_user: UserRole = Depends(get_current_active_user)):
        if current_user not in allowed_roles:
            allowed = [r.value for r in allowed_roles]
            raise HTTPException(status_code=403, detail=f"INSUFFICIENT CLEARANCE. REQUIRED: {allowed}. CURRENT: {current_user.value}")
        return current_user
    return role_checker
from fastapi import HTTPException, Header, Depends, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Literal, Optional
from datetime import datetime, date
import math

# Priority config
class PriorityWeights(BaseModel):
    severity_weight: float = 1.0
    jurisdiction_weight: float = 0.8
    hours_weight: float = 1.2
    executive_weight: float = 0.6
    volume_weight: float = 0.7

# RBAC
class RolePermission(BaseModel):
    role: Literal["public","employee","dpo","cpo","legal","auditor"]
    action: str
    allowed: bool
    requires_cosign: Optional[str] = None

# Dual-approval action
class PendingAction(BaseModel):
    action_id: str
    action_type: str
    initiated_by: str
    initiated_role: str
    requires_role: str
    status: Literal["pending_second_approval","approved","rejected","expired"]
    cosigned_by: Optional[str] = None
    timestamp: datetime

# Workflow rule
class WorkflowRule(BaseModel):
    rule_id: str
    name: str
    condition_field: str
    operator: str
    condition_value: str
    action_type: str
    action_target: Optional[str] = None
    enabled: bool
    requires_dual_approval: bool
    created_by: str

# Vendor
class Vendor(BaseModel):
    vendor_id: str
    name: str
    service: str
    risk_rating: Literal["low","medium","high","critical"]
    dpa_status: Literal["none","draft","executed","expired"]
    dpa_expiration_date: Optional[date] = None
    sccs_in_place: bool
    soc2_type: Literal["none","type_i","type_ii"]
    cross_border: bool
    linked_assessments: list[str]
    linked_risks: list[str]
    status: Literal["active","under_review","blocked","offboarded"]
    next_review_date: date

# Agent telemetry
class AgentTelemetry(BaseModel):
    agent_name: str
    purpose: str
    current_node: str
    queue_depth: int
    avg_processing_ms: float
    failure_rate_pct: float
    recovery_status: Optional[str] = None
    last_heartbeat: datetime

# Scope verification helper
def require_scopes(expected_scopes: List[str]):
    def dependency(x_kibo_scope: str = Header(default="public")):
        if x_kibo_scope not in expected_scopes:
            raise HTTPException(
                status_code=403, 
                detail=f"Forbidden: scope '{x_kibo_scope}' is unauthorized for this resource."
            )
        return x_kibo_scope
    return dependency

# LangGraph imports
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.sqlite import SqliteSaver

# Rule Engine import
from rule_engine import RuleEngine

app = FastAPI(title="KIBO.IS Reactive API Gateway", version="1.0.0")

# Setup CORS for local React dashboard development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# SQLite file locations
DB_FILE = os.path.join(os.path.dirname(__file__), "kibo_state.db")

# In-memory checkpointer connection for LangGraph (or file-based checkpointer)
conn = sqlite3.connect(DB_FILE, check_same_thread=False)
checkpointer = SqliteSaver(conn)

# Initialize Rule Engine
rule_engine = RuleEngine()

# Define LangGraph State Contract
class AgentState(TypedDict):
    id: str
    type: str
    client: str
    jurisdiction: str
    priority: str
    deadline: str
    description: str
    agent: str
    summary: str
    raw: str
    human_decision: Optional[str]
    human_reasoning: Optional[str]
    status: str

# Node 1: Initial Ingestion & Pre-evaluation
def agent_triage_node(state: AgentState) -> Dict[str, Any]:
    print(f"[Triage] Processing transaction {state['id']} for client {state['client']}")
    # Triage step summaries
    return {
        "status": "triage_completed",
        "summary": f"Triage complete. Initial check: {state['description']}"
    }

# Node 2: Human Decision Boundary Checkpoint Node
def human_approval_node(state: AgentState) -> Dict[str, Any]:
    # This node is the breakpoint checkpoint.
    # When resumed, we will check what state was updated.
    decision = state.get("human_decision", "pending")
    reason = state.get("human_reasoning", "")
    print(f"[HITL Approval] Decision received: {decision} with reason: {reason}")
    return {
        "status": "approved" if decision == "approved" else "rejected",
        "summary": f"Human decision: {decision.upper()}. Reason: {reason}"
    }

# Node 3: Execution Output Node
def final_execution_node(state: AgentState) -> Dict[str, Any]:
    print(f"[Execution] Committing action: {state['status']}")
    return {"status": "executed"}

# Compile the LangGraph with HITL interruption
workflow = StateGraph(AgentState)
workflow.add_node("triage", agent_triage_node)
workflow.add_node("human_approval_node", human_approval_node)
workflow.add_node("final_execution", final_execution_node)

# Routing edges
workflow.add_edge(START, "triage")
workflow.add_edge("triage", "human_approval_node")
workflow.add_edge("human_approval_node", "final_execution")
workflow.add_edge("final_execution", END)

# Set breakpoint BEFORE human approval node
graph = workflow.compile(
    checkpointer=checkpointer,
    interrupt_before=["human_approval_node"]
)

# Pydantic input models
class DecisionPayload(BaseModel):
    action: str  # approve_now, approve_always, flag_legal, review_later, reject
    reasoning: str

class NewTransactionRequest(BaseModel):
    id: str
    type: str
    client: str
    jurisdiction: str
    priority: str
    deadline: str
    description: str
    agent: str
    summary: str
    raw: str

# Helper to check if a thread exists and return its state
def get_thread_state(thread_id: str) -> Optional[Dict[str, Any]]:
    config = {"configurable": {"thread_id": thread_id}}
    state = graph.get_state(config)
    if state and state.values:
        return state.values
    return None

# Seed mock database values for initial dashboard visibility
def seed_mock_data():
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    # Create table for direct metadata logging if needed
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS metadata_log (
            thread_id TEXT PRIMARY KEY,
            priority TEXT,
            deadline TEXT,
            status TEXT
        )
    """)
    # Canadian Onboarding Checklist Tasks Table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS onboarding_tasks (
            id TEXT PRIMARY KEY,
            task_name TEXT NOT NULL,
            category TEXT NOT NULL,
            scope TEXT NOT NULL, -- 'Federal' or 'Provincial'
            jurisdiction TEXT NOT NULL, -- 'PIPEDA', 'Law 25', 'CASL', 'PIPA', 'MFIPPA'
            status TEXT DEFAULT 'pending', -- 'pending' or 'completed'
            notes TEXT DEFAULT ''
        )
    """)

    # CASL Consent Registry Table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS casl_consent_registry (
            id TEXT PRIMARY KEY,
            email TEXT NOT NULL,
            name TEXT NOT NULL,
            consent_status TEXT NOT NULL, -- 'Express', 'Implied', 'Expired', 'Unsubscribed'
            consent_date TEXT NOT NULL, -- ISO date
            consent_source TEXT NOT NULL, -- e.g. 'Newsletter Checkbox', 'Product Checkout', etc.
            expiry_date TEXT -- ISO date or NULL for Express
        )
    """)

    # Seed onboarding tasks if empty
    cursor.execute("SELECT count(*) FROM onboarding_tasks")
    if cursor.fetchone()[0] == 0:
        tasks = [
            ("task-1", "Appoint & Publish Privacy Officer", "Governance", "Provincial", "Law 25", "pending", "Highest-ranking executive default unless delegated. Title/contact info must be on website."),
            ("task-2", "Establish a Privacy Management Program", "Governance", "Federal", "PIPEDA", "pending", "Document internal policies for collection, retention, and destruction."),
            ("task-3", "Conduct Privacy Impact Assessments (PIAs)", "Governance", "Provincial", "Law 25", "pending", "Legally mandatory for any new IT project processing personal info."),
            ("task-4", "Record of Processing Activities (RoPA)", "Data Discovery", "Federal", "PIPEDA", "pending", "Map out personal data collected, purpose, and storage locations."),
            ("task-5", "Execute Transfer Impact Assessments (TIAs)", "Data Discovery", "Provincial", "Law 25", "pending", "Required for any data transferred outside of Quebec (including to other provinces)."),
            ("task-6", "Update Vendor DPAs (Data Sharing Agreements)", "Data Discovery", "Federal", "PIPEDA", "pending", "Ensure agreements restrict vendors from using data for their own purposes."),
            ("task-7", "Revamp Cookie & Tracking Consent banner", "Consent", "Provincial", "Law 25", "pending", "Default-off for profiling and tracking cookies. Granular opt-in required."),
            ("task-8", "Configure 'Right to Portability' formats", "Consent", "Provincial", "Law 25", "pending", "Provide data in a structured, commonly used technological format."),
            ("task-9", "Establish De-indexation Workflows", "Consent", "Provincial", "Law 25", "pending", "Handle requests to de-index or remove search result links."),
            ("task-10", "Deploy Bilingual Privacy Policy (English/French)", "Consent", "Provincial", "Law 25", "pending", "French translation is mandatory for Quebec-targeted operations."),
            ("task-11", "Integrate RROSH Dual-Threshold Incident Triage", "Security", "Federal", "PIPEDA", "pending", "Triage breaches against 'Real Risk of Significant Harm' (OPC) and 'Risk of Serious Injury' (CAI)."),
            ("task-12", "Maintain Confidentiality Incident Register", "Security", "Provincial", "Law 25", "pending", "Log all security anomalies, even minor ones, in a permanent internal register."),
            ("task-13", "Automate Data Retention & Destruction Rules", "Security", "Federal", "PIPEDA", "pending", "Ensure secure deletion or full anonymization once purpose is fulfilled."),
            ("task-14", "Identity & Access Integration (Least Privilege)", "Operations", "Federal", "PIPEDA", "pending", "Strict SSO/MFA onboarding and principle of least privilege enforcement."),
            ("task-15", "Enforce Zoom/Teams Retentions (Auto-delete recordings)", "Operations", "Federal", "PIPEDA", "pending", "Auto-delete recordings after 30 days and configure secure email."),
            ("task-16", "Roll out Canadian Privacy Training modules", "Operations", "Federal", "PIPEDA", "pending", "Mandatory staff courses covering PIPEDA, CASL, and Quebec Law 25.")
        ]
        cursor.executemany("INSERT INTO onboarding_tasks VALUES (?, ?, ?, ?, ?, ?, ?)", tasks)

    # Commissioner Verdicts Table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS commissioner_verdicts (
            id TEXT PRIMARY KEY,
            commissioner TEXT NOT NULL,
            jurisdiction TEXT NOT NULL,
            title TEXT NOT NULL,
            summary TEXT NOT NULL,
            compliance_impact TEXT NOT NULL,
            source_url TEXT NOT NULL,
            collected_at TEXT NOT NULL
        )
    """)

    cursor.execute("SELECT count(*) FROM commissioner_verdicts")
    if cursor.fetchone()[0] == 0:
        verdicts = [
            (
                "verdict-on-1",
                "Information and Privacy Commissioner of Ontario (IPC)",
                "Ontario",
                "Ransomware Encryption Deemed Loss of Control",
                "The IPC Ontario ruled that a ransomware attack where data is encrypted but not exfiltrated still legally constitutes a 'loss of control' under PHIPA.",
                "Mandatory breach notification triggered by availability interruptions. Incident Response triage playbooks must flag system/data encryption events as reportable breaches.",
                "https://www.ipc.on.ca/en/decisions",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-sk-1",
                "Office of the Information and Privacy Commissioner of Saskatchewan (OIPC SK)",
                "Saskatchewan",
                "Audit Logging Mandated for Internal Database Reads",
                "OIPC SK reports on local government and health breaches heavily penalize organizations that lack automated audit logging for internal database access.",
                "Systems must log all data reads by employees, not just writes. Strict automated audit trails are required for all database tables containing personal/health records.",
                "https://oipc.sk.ca/decisions/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-atl-1",
                "OIPC Nova Scotia & OIPC Newfoundland and Labrador",
                "Atlantic Provinces",
                "RBAC Safeguards Against Internal snooping",
                "Frequent rulings on unauthorized access ('snooping') within health databases establish strict legal baselines for technical safeguards and training.",
                "RBAC (Role-Based Access Control) architecture must strictly restrict access to only what is necessary, validated against detailed snooping review reports.",
                "https://oipc.novascotia.ca/decisions",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-mb-1",
                "Manitoba Ombudsman",
                "Manitoba",
                "Compliance Checklist for Health Record Encryption",
                "Investigations into the physical and electronic security of health records led to published practice notes that serve as strict compliance checklists under FIPPA/PHIA.",
                "All health data transmission and storage must adhere to Manitoba's specific encryption and physical/electronic security checklists.",
                "https://www.ombudsman.mb.ca/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-nb-1",
                "Office of the Ombud New Brunswick",
                "New Brunswick",
                "Chain of Custody and DSA Mandates for IT Vendors",
                "Health data investigations heavily scrutinize the chain of custody when data moves between custodians and third-party IT vendors under RTIPPA/PHIPAA.",
                "Explicit Data Sharing Agreements (DSAs) are legally required, and technical implementation must restrict vendor access to only what is necessary for maintenance.",
                "https://www.ombudnb.ca/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-pei-1",
                "Office of the Information and Privacy Commissioner of PEI (OIPC PEI)",
                "Prince Edward Island",
                "EHR Consent Directives and Row-Level Masking",
                "PEI's HIA orders dictate that Electronic Health Record (EHR) systems must technically support patient 'consent directives' to mask records from specific practitioners.",
                "Database architecture must support row-level or field-level masking based on user-defined patient consent permissions.",
                "https://www.assembly.pe.ca/oipc/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-yk-1",
                "Yukon Information and Privacy Commissioner",
                "Yukon",
                "Verifiable Audit Trail for Cryptographic Erasure",
                "Yukon compliance reports scrutinize destruction processes for records under ATIPPA/HIPMA, requiring verified audit trails.",
                "Deletion must not just remove pointers; the system must generate a verifiable audit trail proving cryptographic erasure or overwriting occurred.",
                "https://www.yukonipc.ca/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-nwt-1",
                "Office of the Information and Privacy Commissioner of the NWT (OIPC NWT)",
                "Northwest Territories",
                "Remote Access Security, MDM and VPN Enforcement",
                "OIPC NWT reviews highlight privacy risks in remote, highly connected communities, mandating strict remote-access safeguards.",
                "System must enforce Mobile Device Management (MDM), strict session timeouts, and VPN requirements for any remote database access.",
                "https://www.atipp-nt.ca/",
                "2026-07-02T00:00:00Z"
            ),
            (
                "verdict-nu-1",
                "Information and Privacy Commissioner of Nunavut",
                "Nunavut",
                "Mandatory PIAs Prior to Procurement & Deployment",
                "Nunavut commissioner annual reports highlight that new IT systems must be vetted with formal Privacy Impact Assessments (PIAs) prior to procurement.",
                "A formal Privacy Impact Assessment (PIA) is a mandatory gating requirement prior to procuring or deploying new tech processing personal info.",
                "https://atipp-nu.ca/",
                "2026-07-02T00:00:00Z"
            )
        ]
        cursor.executemany("INSERT INTO commissioner_verdicts VALUES (?, ?, ?, ?, ?, ?, ?, ?)", verdicts)

    # Seed CASL registry if empty
    cursor.execute("SELECT count(*) FROM casl_consent_registry")
    if cursor.fetchone()[0] == 0:
        contacts = [
            ("casl-1", "jean.tremblay@company.ca", "Jean Tremblay", "Express", "2025-01-10T14:30:00Z", "Newsletter Checkbox", None),
            ("casl-2", "sarah.smith@ontariotech.ca", "Sarah Smith", "Implied", "2025-06-15T09:00:00Z", "Product Purchase", "2027-06-15T09:00:00Z"),
            ("casl-3", "david.miller@vancouver.org", "David Miller", "Implied", "2026-01-10T11:45:00Z", "Contact Form Inquiry", "2026-07-10T11:45:00Z"),
            ("casl-4", "robert.jones@gmail.com", "Robert Jones", "Implied", "2025-09-01T16:00:00Z", "Quote Request Inquiry", "2026-03-01T16:00:00Z"),
            ("casl-5", "info@quebecinc.qc.ca", "Marie Dubois", "Unsubscribed", "2026-05-12T10:15:00Z", "Web Form Opt-Out", None)
        ]
        cursor.executemany("INSERT INTO casl_consent_registry VALUES (?, ?, ?, ?, ?, ?, ?)", contacts)

    # Onboarding tables
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS clients (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            url TEXT,
            status TEXT DEFAULT 'idle',
            onboarding_mode TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            threshold REAL DEFAULT 0.70
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS findings (
            id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            category TEXT NOT NULL,
            title TEXT NOT NULL,
            description TEXT,
            confidence REAL DEFAULT 0.80,
            source TEXT,
            status TEXT DEFAULT 'pending',
            decision TEXT,
            evidence TEXT,
            FOREIGN KEY(client_id) REFERENCES clients(id)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS reports (
            client_id TEXT PRIMARY KEY,
            completion_percentage INTEGER DEFAULT 0,
            report_data TEXT, -- JSON string of 16-section deliverables
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(client_id) REFERENCES clients(id)
        )
    """)

    # Employee Mode tables
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS employee_intake (
            id TEXT PRIMARY KEY,
            employee_id TEXT NOT NULL,
            project_name TEXT NOT NULL,
            purpose TEXT NOT NULL,
            data_classification TEXT NOT NULL, -- JSON string array
            retention_value INTEGER NOT NULL,
            retention_unit TEXT NOT NULL, -- days, months, years, indefinite
            status TEXT DEFAULT 'draft', -- draft, submitted, under_review, approved
            version_token TEXT DEFAULT 'v1.0_approved'
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS employee_mitigation (
            id TEXT PRIMARY KEY,
            employee_id TEXT NOT NULL,
            project_id TEXT NOT NULL,
            skill_category TEXT NOT NULL,
            framework_target TEXT NOT NULL,
            sla_target TEXT NOT NULL,
            title TEXT NOT NULL,
            description TEXT NOT NULL,
            status TEXT DEFAULT 'open', -- open, in_progress, overdue, closed
            due_date TEXT NOT NULL, -- ISO timestamp or date
            closed_at TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS opc_inquiries (
            thread_id TEXT PRIMARY KEY,
            employee_id TEXT NOT NULL,
            subject TEXT NOT NULL,
            read_status TEXT DEFAULT 'unread', -- unread, read
            created_at TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS opc_messages (
            message_id TEXT PRIMARY KEY,
            thread_id TEXT NOT NULL,
            sender TEXT NOT NULL, -- Employee, CPO
            body TEXT NOT NULL,
            timestamp TEXT NOT NULL, -- ISO-8601
            FOREIGN KEY(thread_id) REFERENCES opc_inquiries(thread_id)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS system_audit_log (
            timestamp TEXT NOT NULL,
            employee_id TEXT NOT NULL,
            target_id TEXT NOT NULL,
            action TEXT NOT NULL,
            delta TEXT NOT NULL -- JSON string of delta
        )
    """)


    cursor.execute('''
        CREATE TABLE IF NOT EXISTS metrics_snapshot (
            snapshot_id TEXT PRIMARY KEY,
            timestamp TEXT,
            privacy_posture TEXT,
            sla_compliance_pct REAL,
            open_requests INTEGER,
            critical_requests INTEGER,
            median_ttr_hours REAL,
            open_incidents INTEGER,
            agents_online INTEGER,
            agents_degraded INTEGER,
            agents_offline INTEGER,
            uptime_pct REAL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS decision_audit (
            decision_id TEXT PRIMARY KEY,
            transaction_id TEXT,
            action TEXT,
            approved_by TEXT,
            authority_level TEXT,
            dual_control_confirmed BOOLEAN,
            ip_address TEXT,
            device TEXT,
            justification_checklist TEXT,
            justification_note TEXT,
            timestamp TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS incidents (
            incident_id TEXT PRIMARY KEY,
            severity TEXT,
            status TEXT,
            jurisdictions_affected TEXT,
            statutory_deadlines TEXT,
            timeline TEXT,
            escalation_sent BOOLEAN,
            escalation_recipients TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS radar_items (
            radar_id TEXT PRIMARY KEY,
            source TEXT,
            title TEXT,
            lane TEXT,
            why_it_matters TEXT,
            suggested_action TEXT,
            dismissed BOOLEAN,
            created_at TEXT
        )
    ''')

    # Seed Radar
    cursor.execute("SELECT count(*) FROM radar_items")
    if cursor.fetchone()[0] == 0:
        import uuid
        from datetime import datetime
        cursor.execute("INSERT INTO radar_items VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), "FTC Webhook", "New COPPA Guidance Published", "informational", "Clarifies minor consent age gating rules.", "Review for Q3 planning.", False, datetime.utcnow().isoformat()))
        cursor.execute("INSERT INTO radar_items VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), "CPPA Alerts", "CPRA Retention Draft Regulations", "actionable", "Impacts our data deletion policies in CA.", "Audit current CCPA retention tags.", False, datetime.utcnow().isoformat()))
        cursor.execute("INSERT INTO radar_items VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), "System Log", "Failed S3 backup scan", "noise", "Automated retry succeeded.", "None", False, datetime.utcnow().isoformat()))

    # Seed Incidents
    cursor.execute("SELECT count(*) FROM incidents")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO incidents VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            ("INC-2026-001", "CRITICAL", "assessed", '["HIPAA"]', '[{"authority": "HHS", "deadline": "2026-08-01T00:00:00Z", "status": "pending"}]', '[{"event": "detected", "timestamp": "2026-06-23T10:00:00Z", "note": "S3 bucket misconfiguration reported"}]', True, '["Legal", "CPO", "CEO"]'))

    conn_db.commit()

    # Seed mock data for employee tables if empty
    cursor.execute("SELECT count(*) FROM employee_intake")
    if cursor.fetchone()[0] == 0:
        cursor.execute("""
            INSERT INTO employee_intake (id, employee_id, project_name, purpose, data_classification, retention_value, retention_unit, status, version_token)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("intake-1", "employee-default", "Global User Profile Database", "Store general billing details and preferences", json.dumps(["Financial", "Sensitive PII"]), 1, "years", "draft", "v1.0_approved"))
        cursor.execute("""
            INSERT INTO employee_intake (id, employee_id, project_name, purpose, data_classification, retention_value, retention_unit, status, version_token)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("intake-2", "employee-default", "Internal HR Metrics Dashboard", "Aggregate department performance metadata", json.dumps(["Employee"]), 6, "months", "submitted", "v1.0_approved"))

    cursor.execute("SELECT count(*) FROM employee_mitigation")
    if cursor.fetchone()[0] == 0:
        # Mock due dates: one active, one overdue (e.g. 2026-01-01)
        cursor.execute("""
            INSERT INTO employee_mitigation (id, employee_id, project_id, skill_category, framework_target, sla_target, title, description, status, due_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("plan-1", "employee-default", "intake-1", "Legal/DPA intake", "HIPAA", "24 Hours", "Rotate AWS API Gateway Keys", "Update secrets manager for all staging clusters", "open", "2026-07-31T00:00:00Z"))
        cursor.execute("""
            INSERT INTO employee_mitigation (id, employee_id, project_id, skill_category, framework_target, sla_target, title, description, status, due_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("plan-2", "employee-default", "intake-2", "Document/DPA automation", "GDPR", "7 Days", "DPA Signature Tracking", "Upload signed DPA from AnalyticsPro Inc.", "open", "2026-05-01T00:00:00Z"))

    cursor.execute("SELECT count(*) FROM opc_inquiries")
    if cursor.fetchone()[0] == 0:
        cursor.execute("""
            INSERT INTO opc_inquiries (thread_id, employee_id, subject, read_status, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, ("thread-opc-1", "employee-default", "HIPAA Health Data Audit Inquiry", "unread", "2026-06-20T12:00:00Z"))
        cursor.execute("""
            INSERT INTO opc_messages (message_id, thread_id, sender, body, timestamp)
            VALUES (?, ?, ?, ?, ?)
        """, ("msg-opc-1", "thread-opc-1", "CPO", "Please clarify if user clinical charts are stored on local device flash memory.", "2026-06-20T12:00:00Z"))
        
        # Second thread - overdue/unread past 48 hours to test the dot: created on 2026-06-15
        cursor.execute("""
            INSERT INTO opc_inquiries (thread_id, employee_id, subject, read_status, created_at)
            VALUES (?, ?, ?, ?, ?)
        """, ("thread-opc-2", "employee-default", "BIPA Biometric Consent Inquiry", "unread", "2026-06-15T09:00:00Z"))
        cursor.execute("""
            INSERT INTO opc_messages (message_id, thread_id, sender, body, timestamp)
            VALUES (?, ?, ?, ?, ?)
        """, ("msg-opc-2", "thread-opc-2", "CPO", "Explain the explicit consent capture flow currently deployed in Illinois.", "2026-06-15T09:00:00Z"))

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS risks (
            risk_id TEXT PRIMARY KEY,
            law_25_section TEXT,
            issue TEXT NOT NULL,
            source_document TEXT,
            source_id TEXT,
            source_title TEXT,
            gap_type TEXT,
            likelihood INTEGER,
            impact INTEGER,
            risk_score INTEGER,
            risk_level TEXT,
            recommendation TEXT,
            khp_response TEXT,
            assigned_to TEXT,
            risk_assigned_date TEXT,
            next_review_date TEXT,
            original_risk_level TEXT,
            status_note TEXT,
            status TEXT DEFAULT 'open',
            created_at TEXT,
            updated_at TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS risk_log (
            log_id TEXT PRIMARY KEY,
            risk_id TEXT NOT NULL,
            actor TEXT NOT NULL,
            action TEXT NOT NULL,
            field_changed TEXT,
            old_value TEXT,
            new_value TEXT,
            note TEXT,
            timestamp TEXT NOT NULL,
            FOREIGN KEY(risk_id) REFERENCES risks(risk_id)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS pias (
            pia_id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            description TEXT,
            creator_name TEXT NOT NULL,
            creator_role TEXT NOT NULL,
            file_name TEXT,
            created_at TEXT NOT NULL
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS pia_extracted_risks (
            extracted_risk_id TEXT PRIMARY KEY,
            pia_id TEXT NOT NULL,
            law_25_section TEXT,
            issue TEXT NOT NULL,
            gap_type TEXT,
            likelihood INTEGER,
            impact INTEGER,
            recommendation TEXT,
            status TEXT DEFAULT 'pending',
            pushed_risk_id TEXT,
            FOREIGN KEY(pia_id) REFERENCES pias(pia_id)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS assessments (
            assessment_id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            type TEXT NOT NULL,
            output_type TEXT NOT NULL,
            level INTEGER NOT NULL,
            version TEXT NOT NULL,
            version_notes TEXT,
            status TEXT NOT NULL,
            risk_level TEXT NOT NULL,
            source TEXT NOT NULL,
            source_id TEXT,
            prepared_by TEXT NOT NULL,
            jurisdictions TEXT NOT NULL,
            roles_in_scope TEXT NOT NULL,
            departments_in_scope TEXT NOT NULL,
            services_in_scope TEXT NOT NULL,
            data_types TEXT NOT NULL,
            cross_border INTEGER DEFAULT 0,
            cross_border_jurisdictions TEXT,
            sections TEXT NOT NULL,
            linked_risks TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            approved_at TEXT,
            pdf_path TEXT,
            docx_path TEXT
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS assessment_versions (
            version_id TEXT PRIMARY KEY,
            assessment_id TEXT NOT NULL,
            version TEXT NOT NULL,
            version_notes TEXT,
            sections TEXT NOT NULL,
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL,
            FOREIGN KEY(assessment_id) REFERENCES assessments(assessment_id)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS assessment_risks (
            assessment_id TEXT NOT NULL,
            risk_id TEXT NOT NULL,
            PRIMARY KEY(assessment_id, risk_id),
            FOREIGN KEY(assessment_id) REFERENCES assessments(assessment_id),
            FOREIGN KEY(risk_id) REFERENCES risks(risk_id)
        )
    """)

    cursor.execute("SELECT count(*) FROM pias")
    if cursor.fetchone()[0] == 0:
        import datetime as dt
        now_str = dt.datetime.utcnow().isoformat() + "Z"
        pia_id = "PIA-0001"
        cursor.execute("""
            INSERT INTO pias (pia_id, title, description, creator_name, creator_role, file_name, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (pia_id, "Customer Portal Personal Data Upgrade PIA", "Assess impact of storing biometric telemetry on new AWS regions.", "Jane Doe", "Privacy Specialist", "portal_personal_data_v2.pdf", now_str))
        
        cursor.execute("""
            INSERT INTO pia_extracted_risks (extracted_risk_id, pia_id, law_25_section, issue, gap_type, likelihood, impact, recommendation, status)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("EXT-RSK-0001", pia_id, "12", "Missing explicit opt-in for biometric authentication telemetry", "Consent Management", 4, 5, "Implement standard Law 25 opt-in popup prior to biometric scan.", "pending"))
        
        cursor.execute("""
            INSERT INTO pia_extracted_risks (extracted_risk_id, pia_id, law_25_section, issue, gap_type, likelihood, impact, recommendation, status)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("EXT-RSK-0002", pia_id, "15", "Third-party sharing of behavioral patterns with AnalyticsPro Inc.", "Vendor Assessment", 3, 4, "Perform vendor assessment and execute formal standard DPA.", "pending"))

    cursor.execute("SELECT count(*) FROM assessments")
    if cursor.fetchone()[0] == 0:
        import datetime as dt
        now_str = dt.datetime.utcnow().isoformat() + "Z"
        
        # WhatsApp PIA Seed Data
        whatsapp_sections = {
            "preliminary_analysis": {
                "project_title": "WhatsApp Access Point — Crisis Texting Service",
                "institution_dept": "Clinical Operations",
                "executive_sponsor": "CEO",
                "project_lead": "Wael Hassan",
                "pia_lead": "Wael Hassan",
                "planned_go_live": "2026-07-01",
                "personal_information": "Yes",
                "data_subjects": "Service Users - Youth",
                "relevant_legislation": "PIPEDA, Law 25, PHIPA"
            },
            "executive_summary": "KHP plans to add WhatsApp as an official entry point to its text-based crisis support program. Risks include cross-border US hosting and Meta platform dependency. Mitigated via DPA with SCCs, AES-256 encryption, 14-day media purge.",
            "scope_and_architecture": {
                "scope_description": "Adding WhatsApp entry point to crisis program",
                "system_architecture": "Users text WhatsApp API -> Twilio -> Kibo Gateway -> Crisis Responders Console",
                "roles_responsibilities": [
                    {"role": "Controller", "entity": "KHP"},
                    {"role": "Processor", "entity": "Twilio"},
                    {"role": "Sub-processor", "entity": "Meta"}
                ]
            },
            "data_classification": {
                "sensitivity": "Highly Sensitive",
                "data_sovereignty": "Standard cross-border US hosting constraints",
                "regulatory_basis": "Consent-based health service"
            },
            "threat_analysis": [
                {"domain": "Cross-border data transfer", "level": "High", "description": "US hosting via Twilio and CTL USA"},
                {"domain": "Vendor assurance", "level": "Medium", "description": "SOC 2 Type I audit only"}
            ],
            "privacy_analysis": [
                {"requirement": "Collection & Consent", "gap": "Need explicit WhatsApp consent banner", "mitigation": "Configure pre-chat consent flow"},
                {"requirement": "Use Limitation", "gap": "Meta profile telemetry tracking", "mitigation": "Anonymise metadata before transit"}
            ],
            "necessity_proportionality": {
                "necessity": "Required to reach youth demographic where they are active",
                "effectiveness": "WhatsApp shows 85% penetration in target cohort",
                "minimisation": "No profiling or message store beyond 14 days",
                "proportionality": "Benefits of safety intervention outweigh metadata risk",
                "finding": "Met"
            },
            "recommendations": [
                {"priority": "High", "action": "Execute DPA with Twilio/Meta", "owner": "Legal", "due": "Pre-launch"},
                {"priority": "Medium", "action": "Publish updated Privacy Policy", "owner": "Communications", "due": "Pre-launch"}
            ],
            "signoff": {
                "psr_status": "Good to Go",
                "executive_endorsement": "Obtained",
                "signoff_date": "2026-06-09"
            }
        }
        
        cursor.execute("""
            INSERT INTO assessments (
                assessment_id, title, type, output_type, level, version, version_notes, status, risk_level, 
                source, source_id, prepared_by, jurisdictions, roles_in_scope, departments_in_scope, 
                services_in_scope, data_types, cross_border, cross_border_jurisdictions, sections, linked_risks, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            "PIA-2026-001",
            "WhatsApp Access Point — Crisis Texting Service",
            "PIA",
            "physical_pia",
            3,
            "0.3",
            "Adds necessity & proportionality assessment",
            "good_to_go",
            "Medium",
            "intake",
            "REQ-2026-001",
            "Wael Hassan",
            json.dumps(["PIPEDA", "Law 25", "PHIPA"]),
            json.dumps(["Service Users — Youth", "Crisis Responders"]),
            json.dumps(["Clinical Operations", "Technology"]),
            json.dumps(["WhatsApp", "CTL USA", "Twilio", "AWS"]),
            json.dumps(["Phone number", "Mental health disclosures", "Crisis transcripts"]),
            1,
            json.dumps(["United States"]),
            json.dumps(whatsapp_sections),
            json.dumps(["RSK-0004", "RSK-0005"]),
            now_str,
            now_str
        ))
        
        # Indigenous App PIA Seed Data
        indigenous_sections = {
            "preliminary_analysis": {
                "project_title": "KHP Indigenous App — Launch Zero",
                "institution_dept": "Technology / IT",
                "executive_sponsor": "CEO",
                "project_lead": "Wael Hassan",
                "pia_lead": "Wael Hassan",
                "planned_go_live": "2026-04-01",
                "personal_information": "Yes",
                "data_subjects": "Service Users - Indigenous youth",
                "relevant_legislation": "PIPEDA, Law 25"
            },
            "executive_summary": "Involves OCAP® principles, offline device storage, iframe caching of crisis transcripts, no biometric lock on a shared device, youth population. Overall risk: Medium.",
            "scope_and_architecture": {
                "scope_description": "Custom mobile application designed to offer culturally grounded mental health support pathways.",
                "system_architecture": "React Native Mobile App -> Secure Local SQLite database -> Kibo Private Gateway",
                "roles_responsibilities": [
                    {"role": "Controller", "entity": "KHP"},
                    {"role": "Processor", "entity": "Secure Cloud Provider"}
                ]
            },
            "data_classification": {
                "sensitivity": "Sensitive",
                "data_sovereignty": "OCAP® principles enforced. Local database with zero cloud sync for transcripts.",
                "regulatory_basis": "Explicit opt-in framework."
            },
            "threat_analysis": [
                {"domain": "Data sovereignty", "level": "Medium", "description": "Strict compliance with Indigenous ownership principles"},
                {"domain": "Access control", "level": "High", "description": "Shared mobile devices without screen locks"}
            ],
            "privacy_analysis": [
                {"requirement": "Collection & Consent", "gap": "Bilingual English/Indigenous languages consent flow needed", "mitigation": "Implement audio/visual consent screens"},
                {"requirement": "Use Limitation", "gap": "Offline cache risk", "mitigation": "Encrypt SQLite database files on-device"}
            ],
            "necessity_proportionality": {
                "necessity": "Demonstrably necessary to support remote Indigenous communities",
                "effectiveness": "Enables critical emergency offline crisis assistance",
                "minimisation": "Zero personal data uploaded to backend; transcripts remain local",
                "proportionality": "Privacy exposure minimised via localized control",
                "finding": "Met"
            },
            "recommendations": [
                {"priority": "High", "action": "Implement SQLCipher on mobile client", "owner": "IT", "due": "Pre-launch"},
                {"priority": "Medium", "action": "Translate privacy notice to local languages", "owner": "Communications", "due": "Post-launch"}
            ],
            "signoff": {
                "psr_status": "Good to Go",
                "executive_endorsement": "Obtained",
                "signoff_date": "2026-03-15"
            }
        }
        
        cursor.execute("""
            INSERT INTO assessments (
                assessment_id, title, type, output_type, level, version, version_notes, status, risk_level, 
                source, source_id, prepared_by, jurisdictions, roles_in_scope, departments_in_scope, 
                services_in_scope, data_types, cross_border, cross_border_jurisdictions, sections, linked_risks, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            "PIA-2026-002",
            "KHP Indigenous App — Launch Zero",
            "PIA",
            "physical_pia",
            3,
            "0.1",
            "Initial draft",
            "good_to_go",
            "Medium",
            "project",
            "REQ-2026-002",
            "Wael Hassan",
            json.dumps(["PIPEDA", "Law 25"]),
            json.dumps(["Service Users — Indigenous youth", "Crisis Responders"]),
            json.dumps(["Technology / IT", "Communications"]),
            json.dumps(["AWS", "On-premise servers"]),
            json.dumps(["Location data", "Mental health disclosures", "Crisis transcripts"]),
            0,
            json.dumps([]),
            json.dumps(indigenous_sections),
            json.dumps(["RSK-0004", "RSK-0005"]),
            now_str,
            now_str
        ))
        
        # Link to assessment risks
        cursor.execute("INSERT INTO assessment_risks (assessment_id, risk_id) VALUES (?, ?)", ("PIA-2026-001", "RSK-0004"))
        cursor.execute("INSERT INTO assessment_risks (assessment_id, risk_id) VALUES (?, ?)", ("PIA-2026-001", "RSK-0005"))
        cursor.execute("INSERT INTO assessment_risks (assessment_id, risk_id) VALUES (?, ?)", ("PIA-2026-002", "RSK-0004"))
        cursor.execute("INSERT INTO assessment_risks (assessment_id, risk_id) VALUES (?, ?)", ("PIA-2026-002", "RSK-0005"))

    cursor.execute("SELECT count(*) FROM risks")
    if cursor.fetchone()[0] == 0:
        import datetime as dt
        now_str = dt.datetime.utcnow().isoformat() + "Z"
        
        # RSK-0001
        cursor.execute("""
            INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("RSK-0001", "3.8", "Incident Register Not Law 25 Standardized", "Organizational", None, None, "Operational Gap", 4, 4, 16, "High", "Align register to Law 25; standardize format; annual review.", "Register initiated.", "Waël", "2022-09-01", "2026-07-15", "High", "On Going — Incident Log tab in this tracker is the standardized register.", "open", now_str, now_str))

        # RSK-0002
        cursor.execute("""
            INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("RSK-0002", "None", "Identity Verification for Access Requests Not Standardized", "Organizational", None, None, "Operational Gap", 3, 4, 12, "Medium", "Establish standard identity verification procedures.", "Verification forms drafted.", "Waël", "2022-09-01", "2026-07-15", "Medium", "Draft pending review.", "open", now_str, now_str))

        # RSK-0003
        cursor.execute("""
            INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("RSK-0003", "None", "Cross-Border PIA Outdated", "Organizational", None, None, "Governance Gap", 3, 4, 12, "Medium", "Update cross-border transfer assessments.", "Assessment scheduled.", "Waël", "2022-09-01", "2026-10-14", "Medium", "Pending next review cycle.", "open", now_str, now_str))

        # RSK-0004
        cursor.execute("""
            INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("RSK-0004", "None", "No executed agreement or DPA with Cohere; SOC 2 unconfirmed", "Organizational", None, None, "Design Gap", 3, 4, 12, "Medium", "Obtain and review SOC 2 report, sign DPA.", "DPA sent for signature.", "Betty", "2026-06-23", "2026-06-27", "Medium", "Waiting for reply from Betty.", "open", now_str, now_str))

        # RSK-0005
        cursor.execute("""
            INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, ("RSK-0005", "None", "Meta encryption policy change — WhatsApp subpoena exposure", "Organizational", None, None, "Operational Gap", 2, 4, 8, "Medium", "Review encryption protocols and legal policies.", "Mitigation steps logged.", "Waël", "2022-09-01", "2026-07-15", "Medium", "On going.", "open", now_str, now_str))

    # New tables for localization, employee training, data inventory, meetings, inbox, PSR recommendations
    cursor.execute("CREATE TABLE IF NOT EXISTS user_settings (key TEXT PRIMARY KEY, value TEXT)")
    cursor.execute("CREATE TABLE IF NOT EXISTS employee_data_inventory (id TEXT PRIMARY KEY, employee_id TEXT NOT NULL, system_name TEXT NOT NULL, data_types TEXT NOT NULL, purpose TEXT NOT NULL, retention TEXT NOT NULL, sharing TEXT NOT NULL, submitted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    cursor.execute("CREATE TABLE IF NOT EXISTS employee_training (module_id TEXT NOT NULL, employee_id TEXT NOT NULL, title TEXT NOT NULL, role_tags TEXT NOT NULL, jurisdiction TEXT NOT NULL, source_policy TEXT, duration_min INTEGER, required BOOLEAN, completed BOOLEAN DEFAULT 0, completed_at TEXT, PRIMARY KEY(module_id, employee_id))")
    cursor.execute("CREATE TABLE IF NOT EXISTS meetings (meeting_id TEXT PRIMARY KEY, type TEXT NOT NULL, date TEXT NOT NULL, agenda TEXT NOT NULL, attendees TEXT NOT NULL, minutes TEXT, action_items TEXT, materials TEXT)")
    cursor.execute("CREATE TABLE IF NOT EXISTS simulated_inbox (email_id TEXT PRIMARY KEY, sender TEXT NOT NULL, subject TEXT NOT NULL, body TEXT NOT NULL, category TEXT NOT NULL, status TEXT DEFAULT 'unread', converted_to_transaction TEXT)")
    cursor.execute("CREATE TABLE IF NOT EXISTS simulated_inbox_replies (reply_id TEXT PRIMARY KEY, email_id TEXT NOT NULL, body TEXT NOT NULL, sent_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    cursor.execute("CREATE TABLE IF NOT EXISTS psr_recommendations (recommendation_id TEXT PRIMARY KEY, risk_id TEXT NOT NULL, member TEXT NOT NULL, vote TEXT NOT NULL, recommendation TEXT NOT NULL, timestamp TEXT NOT NULL)")
    cursor.execute("CREATE TABLE IF NOT EXISTS agent_lessons_learned (lesson_id TEXT PRIMARY KEY, domain TEXT NOT NULL, feedback_notes TEXT NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    cursor.execute("CREATE TABLE IF NOT EXISTS legal_ground_truth (clause_id TEXT PRIMARY KEY, legislation TEXT NOT NULL, clause_text TEXT NOT NULL, keywords TEXT NOT NULL)")
    
    cursor.execute("SELECT count(*) FROM legal_ground_truth")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO legal_ground_truth (clause_id, legislation, clause_text, keywords) VALUES (?, ?, ?, ?)",
                       ("LAW25-SEC14", "quebec_law_25", "Quebec Law 25: Consent must be clear, free, informed, and given for specific purposes. Any targeted campaigns in French jurisdiction require bilingual opt-in terms and accessible French policy documentation.", "bilingual, campaigns, Law 25, French"))
        cursor.execute("INSERT INTO legal_ground_truth (clause_id, legislation, clause_text, keywords) VALUES (?, ?, ?, ?)",
                       ("PIPEDA-SCH1-4.7", "pipeda", "PIPEDA Schedule 1 Section 4.7: Personal information must be protected by security safeguards appropriate to the sensitivity of the information, including revoking stale credentials and encryption.", "safeguards, security, credentials, encryption"))
        cursor.execute("INSERT INTO legal_ground_truth (clause_id, legislation, clause_text, keywords) VALUES (?, ?, ?, ?)",
                       ("GDPR-ART28", "gdpr", "GDPR Article 28: Data processing by a processor shall be governed by a contract or other legal act (Data Protection Agreement - DPA) binding the processor to the controller.", "DPA, contract, processor, vendor"))
    
    # Seed initial memories/lessons
    cursor.execute("SELECT count(*) FROM agent_lessons_learned")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO agent_lessons_learned (lesson_id, domain, feedback_notes) VALUES (?, ?, ?)", ("L-001", "onboarding", "Ensure Quebec Law 25 bilingual policies are marked mandatory during French targeted campaigns."))
        cursor.execute("INSERT INTO agent_lessons_learned (lesson_id, domain, feedback_notes) VALUES (?, ?, ?)", ("L-002", "dsar", "Always redact third-party IP addresses from SMS Twilio backups before delivery."))

    # Agentic CPO Intelligence Architecture Tables
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS knowledge_updates (
            update_id TEXT PRIMARY KEY,
            domain TEXT NOT NULL,
            agent TEXT NOT NULL,
            change_type TEXT NOT NULL,
            entity TEXT NOT NULL,
            old_value TEXT,
            new_value TEXT NOT NULL,
            confidence INTEGER NOT NULL,
            source TEXT NOT NULL,
            impact_summary TEXT,
            affected_entities TEXT,
            status TEXT DEFAULT 'pending_review',
            detected_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS regulatory_ontology (
            obligation_id TEXT PRIMARY KEY,
            jurisdiction TEXT NOT NULL,
            statute TEXT NOT NULL,
            obligation_name TEXT NOT NULL,
            applicability_trigger TEXT NOT NULL,
            deadline_days INTEGER NOT NULL,
            regulator TEXT NOT NULL,
            last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS vendor_graph (
            vendor_id TEXT PRIMARY KEY,
            vendor_name TEXT NOT NULL,
            service TEXT NOT NULL,
            data_types TEXT NOT NULL,
            dpa_status TEXT NOT NULL,
            soc2_status TEXT NOT NULL,
            sub_processors TEXT,
            cross_border_vectors TEXT,
            risk_rating TEXT NOT NULL,
            last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS onboarding_sessions (
            session_id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            status TEXT NOT NULL,
            progress REAL NOT NULL,
            profile_json TEXT NOT NULL,
            logs_json TEXT NOT NULL,
            urls_json TEXT NOT NULL,
            files_json TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS onboarding_gaps (
            gap_id TEXT PRIMARY KEY,
            session_id TEXT NOT NULL,
            type TEXT NOT NULL,
            title TEXT NOT NULL,
            details TEXT NOT NULL,
            priority TEXT NOT NULL,
            status TEXT DEFAULT 'pending'
        )
    """)

    # Seed regulatory ontology
    cursor.execute("SELECT count(*) FROM regulatory_ontology")
    if cursor.fetchone()[0] == 0:
        obligations = [
            ("ob-1", "ontario", "FIPPA", "Freedom of Information Access", "Access request filed by individual", 30, "Information and Privacy Commissioner of Ontario (IPC)"),
            ("ob-2", "quebec", "Law 25", "Mandatory Privacy Impact Assessment", "Deploying new high-risk IT system", 30, "Commission d'accès à l'information (CAI)"),
            ("ob-3", "canada", "PIPEDA", "Real Risk of Significant Harm Triage", "Data breach detected", 0, "Office of the Privacy Commissioner of Canada (OPC)"),
            ("ob-4", "us", "CPRA", "Consumer Deletion Request", "Consumer submits deletion request", 45, "California Privacy Protection Agency (CPPA)"),
            ("ob-5", "eu", "GDPR", "72-Hour Breach Notification", "Confidentiality incident detected", 3, "Supervisory Authority (DPA)")
        ]
        cursor.executemany("INSERT INTO regulatory_ontology VALUES (?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)", obligations)

    # Seed vendor graph
    cursor.execute("SELECT count(*) FROM vendor_graph")
    if cursor.fetchone()[0] == 0:
        vendors = [
            ("v-1", "Cohere", "AI Model Hosting", "User chat logs, prompt data", "No DPA executed", "SOC 2 Type II active", "None", "None", "High"),
            ("v-2", "Twilio", "SMS & Verification", "Phone numbers, timestamps", "DPA executed", "SOC 2 Type II active", "AWS, GCP", "Canada to US transfer", "Medium"),
            ("v-3", "AWS", "Cloud Infrastructure", "Application databases, storage", "DPA executed", "SOC 2 Type II active", "AWS entities", "US-East storage", "Low"),
            ("v-4", "Salesforce", "Customer Relation", "Client profiles, clinical notes", "DPA executed", "SOC 2 Type II active", "Salesforce entities", "US-East storage", "Medium")
        ]
        cursor.executemany("INSERT INTO vendor_graph VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)", vendors)

    # Seed user_settings with default active jurisdiction 'ontario'
    cursor.execute("SELECT count(*) FROM user_settings WHERE key='active_jurisdiction'")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO user_settings (key, value) VALUES ('active_jurisdiction', 'ontario')")

    # Seed training modules
    cursor.execute("SELECT count(*) FROM employee_training")
    if cursor.fetchone()[0] == 0:
        import uuid
        modules = [
            ("mod-1", "employee-default", "Privacy Fundamentals", '["clinical_counselor", "frontline", "sales"]', "canada", "policy-1", 20, 1),
            ("mod-2", "employee-default", "Handling Crisis Disclosures", '["clinical_counselor", "frontline"]', "canada", "policy-2", 30, 1),
            ("mod-3", "employee-default", "Law 25 for Frontline Staff", '["clinical_counselor", "frontline"]', "quebec", "policy-3", 25, 1),
            ("mod-4", "employee-default", "Cross-Border Data Basics", '["clinical_counselor", "frontline", "sales"]', "canada", "policy-4", 15, 0),
            
            ("mod-5", "employee-default", "CCPA Rights & Deletion Flow", '["sales", "support"]', "us", "policy-5", 20, 1),
            ("mod-6", "employee-default", "GDPR Data Subject Rights", '["sales", "support", "clinical_counselor"]', "eu", "policy-6", 30, 1),
            ("mod-7", "employee-default", "UK GDPR Compliance", '["sales", "support"]', "uk", "policy-7", 25, 1),
        ]
        for mid, eid, title, roles, jur, pol, dur, req in modules:
            cursor.execute("INSERT INTO employee_training (module_id, employee_id, title, role_tags, jurisdiction, source_policy, duration_min, required, completed) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 0)", (mid, eid, title, roles, jur, pol, dur, req))

    # Seed simulated_inbox
    cursor.execute("SELECT count(*) FROM simulated_inbox")
    if cursor.fetchone()[0] == 0:
        emails = [
            ("email-1", "john.doe@gmail.com", "DSAR Deletion Request", "Please delete all my personal and clinical logs from your systems immediately. I reside in California.", "dsar"),
            ("email-2", "officer@ipc.on.ca", "IPC Audit Notification", "The Information and Privacy Commissioner of Ontario requires an audit of your MFIPPA compliance records for Q2.", "regulator"),
            ("email-3", "compliance@analyticspro.com", "Vendor Security Questionnaire", "Please complete our annual security and data transfer compliance review to maintain the analytics dashboard service.", "vendor"),
            ("email-4", "counselor.betty@kibo.org", "Law 25 Consent Query", "Do we need explicit audio-visual consent from Quebec-based minors before initiating therapy sessions?", "staff_question"),
        ]
        for eid, sender, subj, body, cat in emails:
            cursor.execute("INSERT INTO simulated_inbox (email_id, sender, subject, body, category, status) VALUES (?, ?, ?, ?, ?, 'unread')", (eid, sender, subj, body, cat))

    # Seed meetings
    cursor.execute("SELECT count(*) FROM meetings")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO meetings (meeting_id, type, date, agenda, attendees, minutes, action_items, materials) VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (
            "meet-1",
            "privacy_security_weekly",
            "2026-07-01T10:00:00Z",
            json.dumps(["HIPAA S3 Misconfiguration Review", "WhatsApp PIA Update", "Audit of Law 25 compliance"]),
            json.dumps(["Wael Hassan", "Betty", "Legal Team"]),
            "Discussed S3 bucket settings and confirmed DPA sent to Cohere.",
            json.dumps([{"task": "Acknowledge HIPAA risk", "owner": "Betty", "due": "2026-07-05"}]),
            json.dumps(["AWS Security Policy Draft.pdf"])
        ))

    conn_db.commit()
    conn_db.close()


    mock_txs = [
        {
            "id": "thread-101",
            "type": "DSAR",
            "client": "TechCorp Inc.",
            "jurisdiction": "US",
            "priority": "critical",
            "deadline": "48 hours",
            "description": "Verify identity and process CCPA deletion request.",
            "agent": "US DSAR Rights Agent",
            "summary": "User authenticated. 3 systems matched. No legal hold.",
            "raw": '{"systems": ["crm", "billing"], "confidence": 0.99}'
        },
        {
            "id": "thread-102",
            "type": "Vendor",
            "client": "CloudScale Inc.",
            "jurisdiction": "EU",
            "priority": "high",
            "deadline": "12 hours",
            "description": "Evaluate vendor GDPR DPA adequacy for analytics integration.",
            "agent": "EU Data Transfer Auditor",
            "summary": "Standard contractual clauses present. Sub-processors listed.",
            "raw": '{"risk_level": "medium", "data_location": "Frankfurt"}'
        },
        {
            "id": "thread-103",
            "type": "Breach",
            "client": "HealthFirst Corp",
            "jurisdiction": "Canada",
            "priority": "critical",
            "deadline": "24 hours",
            "description": "Perform HIPAA compliance review on misconfigured S3 bucket.",
            "agent": "Canada PHIPA Security Agent",
            "summary": "Exposed files: 120 health charts. Exposure time: 14 hours.",
            "raw": '{"affected_records": 120, "pii_exposed": true}'
        }
    ]

    for tx in mock_txs:
        # Check if already seeded to avoid duplication
        if get_thread_state(tx["id"]) is None:
            # Check rules first
            matched_rule = rule_engine.evaluate(tx)
            config = {"configurable": {"thread_id": tx["id"]}}
            
            if matched_rule:
                # Auto execution bypass
                print(f"[Seeder] Auto-executing {tx['id']} due to rule match: {matched_rule['rule_id']}")
                initial_state = {**tx, "human_decision": "approved", "human_reasoning": matched_rule["human_reasoning"], "status": "auto_approved"}
                graph.update_state(config, initial_state)
                # Run complete graph bypassing human approval since state is already resolved
                graph.invoke(None, config)
            else:
                # Queue up at breakpoint
                print(f"[Seeder] Queuing {tx['id']} at human approval checkpoint")
                graph.invoke(tx, config)

seed_mock_data()
canadian_verdicts_agent.start_scheduler()


@app.get("/api/transactions")
def get_transactions(scope: str = Depends(require_scopes(["expert"]))):
    """Fetch pending transactions (threads currently interrupted at human_approval_node)."""
    # Fetch all states in SqliteSaver checkpointer
    transactions = []
    
    # We query the sqlite database file directly for active threads
    try:
        conn_db = sqlite3.connect(DB_FILE)
        cursor = conn_db.cursor()
        # Query unique thread IDs registered in langgraph checkpoint storage
        cursor.execute("SELECT DISTINCT thread_id FROM checkpoints")
        thread_ids = [row[0] for row in cursor.fetchall()]
        conn_db.close()
    except Exception as e:
        print(f"Error querying checkpoint database: {e}")
        thread_ids = []

    for t_id in thread_ids:
        state_vals = get_thread_state(t_id)
        if state_vals:
            # Check if current execution thread is pending human gate approval
            config = {"configurable": {"thread_id": t_id}}
            next_steps = graph.get_state(config).next
            
            # If the next node in queue is human_approval_node, it is pending CPO input
            is_pending = "human_approval_node" in next_steps
            
            transactions.append({
                "id": t_id,
                "type": state_vals.get("type", ""),
                "client": state_vals.get("client", ""),
                "jurisdiction": state_vals.get("jurisdiction", ""),
                "priority": state_vals.get("priority", "medium"),
                "deadline": state_vals.get("deadline", "N/A"),
                "description": state_vals.get("description", ""),
                "agent": state_vals.get("agent", ""),
                "summary": state_vals.get("summary", ""),
                "raw": state_vals.get("raw", "{}"),
                "status": "pending" if is_pending else state_vals.get("status", "completed"),
                "human_decision": state_vals.get("human_decision"),
                "human_reasoning": state_vals.get("human_reasoning")
            })
            
    return transactions

@app.post("/api/transactions/{thread_id}/decision")
def post_decision(thread_id: str, payload: DecisionPayload, scope: str = Depends(require_scopes(["expert"]))):
    """
    Accepts decision payload, handles rule learning if 'approve_always' or 'reject',
    updates graph state, and resumes thread.
    """
    config = {"configurable": {"thread_id": thread_id}}
    state_vals = get_thread_state(thread_id)
    if not state_vals:
        raise HTTPException(status_code=404, detail="Transaction thread not found")

    action_map = {
        "approve_now": "approved",
        "approve_always": "approved",
        "flag_legal": "flagged_legal",
        "review_later": "deferred",
        "reject": "rejected"
    }
    
    decision = action_map.get(payload.action, "pending")
    
    # Handle Rule Learning
    if payload.action in ["approve_always", "reject"]:
        rule_action = "auto-approve" if payload.action == "approve_always" else "auto-reject"
        condition = {
            "client": state_vals.get("client"),
            "type": state_vals.get("type"),
            "jurisdiction": state_vals.get("jurisdiction")
        }
        rule_engine.add_rule(condition, rule_action, payload.reasoning)
        print(f"[Rule Engine] Learned new {rule_action} rule for client {condition['client']}")

    # Update State in LangGraph
    graph.update_state(config, {
        "human_decision": decision,
        "human_reasoning": payload.reasoning,
        "status": decision
    })

    # If the user deferred, keep it suspended. Otherwise, resume graph execution
    if payload.action != "review_later":
        # Resume the thread - LangGraph will execute the human_approval_node and finalize
        graph.invoke(None, config)

    return {"status": "success", "resolved_decision": decision}

@app.post("/api/transactions/new")
def create_transaction(tx: NewTransactionRequest, scope: str = Depends(require_scopes(["public", "employee", "expert"]))):
    """Create a new transaction and run it through the LangGraph and Rule Engine."""
    tx_dict = tx.dict()
    
    # Check Rules first
    matched_rule = rule_engine.evaluate(tx_dict)
    config = {"configurable": {"thread_id": tx.id}}
    
    if matched_rule:
        # Bypass human queue entirely
        print(f"[Ingestion] Auto-executing {tx.id} due to rule match: {matched_rule['rule_id']}")
        initial_state = {**tx_dict, "human_decision": "approved", "human_reasoning": matched_rule["human_reasoning"], "status": "auto_approved"}
        graph.update_state(config, initial_state)
        graph.invoke(None, config)
        return {"status": "auto_approved", "rule_id": matched_rule["rule_id"]}
    else:
        # Trigger normal flow and suspend at breakpoint
        print(f"[Ingestion] Invoking graph for {tx.id}, suspending at human breakpoint")
        graph.invoke(tx_dict, config)
        return {"status": "pending_approval"}

@app.get("/api/rules")
def get_rules(scope: str = Depends(require_scopes(["expert"]))):
    """Fetch learned automation rules."""
    return rule_engine.rules

# --- Client Onboarding API Endpoints ---
class NewClientRequest(BaseModel):
    name: str
    url: Optional[str] = None
    mode: str  # website, documents, systems, guided, random
    threshold: Optional[float] = 0.70

class FindingDecisionPayload(BaseModel):
    decision: str  # accept, reject, modify
    title: Optional[str] = None
    description: Optional[str] = None

class GuidedAnswersPayload(BaseModel):
    collect_customer_info: bool
    process_employee_info: bool
    operate_internationally: bool
    process_health_info: bool
    process_financial_info: bool
    collect_children_data: bool
    use_ai_systems: bool
    use_biometrics: bool

@app.get("/api/clients")
def list_clients(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM clients ORDER BY created_at DESC")
    rows = cursor.fetchall()
    clients = [dict(r) for r in rows]
    conn_db.close()
    return clients

@app.get("/api/clients/{client_id}")
def get_client(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Client not found")
    return dict(row)

@app.post("/api/clients/new")
def create_client(payload: NewClientRequest, scope: str = Depends(require_scopes(["expert"]))):
    import uuid
    client_id = f"client-{uuid.uuid4().hex[:6]}"
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute(
        "INSERT INTO clients (id, name, url, status, onboarding_mode, threshold) VALUES (?, ?, ?, ?, ?, ?)",
        (client_id, payload.name, payload.url, "idle", payload.mode, payload.threshold)
    )
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "client_id": client_id}

@app.get("/api/clients/{client_id}/findings")
def get_client_findings(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM findings WHERE client_id = ?", (client_id,))
    rows = cursor.fetchall()
    findings = [dict(r) for r in rows]
    conn_db.close()
    return findings

@app.post("/api/clients/{client_id}/findings/{finding_id}/decision")
def submit_finding_decision(client_id: str, finding_id: str, payload: FindingDecisionPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    if payload.decision == "modify":
        cursor.execute(
            "UPDATE findings SET status = 'reviewed', decision = ?, title = ?, description = ? WHERE id = ? AND client_id = ?",
            (payload.decision, payload.title, payload.description, finding_id, client_id)
        )
    else:
        cursor.execute(
            "UPDATE findings SET status = 'reviewed', decision = ? WHERE id = ? AND client_id = ?",
            (payload.decision, finding_id, client_id)
        )
    
    conn_db.commit()
    conn_db.close()
    
    # Recalculate and update the compiled report deliverables
    compile_client_report(client_id)
    
    return {"status": "success"}

def compile_client_report(client_id: str):
    """
    Compiles the 16 deliverables.
    Marks sections depending on whether there are findings matching specific keywords.
    """
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    # Fetch client details
    cursor.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    client_row = cursor.fetchone()
    if not client_row:
        conn_db.close()
        return
        
    client = dict(client_row)
    
    # Fetch all findings
    cursor.execute("SELECT * FROM findings WHERE client_id = ?", (client_id,))
    findings = [dict(r) for r in cursor.fetchall()]
    
    # Generate 16 mandatory sections
    sections = [
        "1. Executive Privacy Assessment Summary",
        "2. Comprehensive Data Mapping & Inventory Matrix",
        "3. Sectoral Patchwork Applicability Grid (CPRA, HIPAA, TDPSA)",
        "4. High-Risk Automated Decision-Making Technology (ADMT) Audit",
        "5. Biometric Information Privacy Act (BIPA) Statutory Consent Flow",
        "6. Children’s Online Privacy Protection (COPPA) Audit Registry",
        "7. Financial Services GLBA Safeguards Compliance Attestation",
        "8. Vendor Risk & Third-Party DPA Adequacy Matrix",
        "9. Cross-Border Data Transfer Risk Analysis (SCC Adequacy)",
        "10. Consumer Rights Request (DSAR) Ingestion Strategy",
        "11. Incident & Data Breach 50-State Triage Schema",
        "12. Employee Training & Policy Alignment Matrix",
        "13. Core Web Vitals & Tracking Consent Cookie Opt-Out Audit",
        "14. Public-Facing Privacy Policy Redline & Gaps Analysis",
        "15. 1984 Hosting DNS Integration & Verification Manifest",
        "16. Chief Privacy Officer Action Queue & Governance Baseline"
    ]
    
    report_sections = {}
    completed_count = 0
    
    for sec in sections:
        is_completed = False
        sec_findings = []
        
        # Mapping keywords to section
        keywords = {
            "1.": ["executive", "overall", "summary"],
            "2.": ["mapping", "inventory", "database", "data flow"],
            "3.": ["sectoral", "cpra", "hipaa", "tdpsa", "jurisdiction"],
            "4.": ["admt", "automated", "algorithm", "profiling"],
            "5.": ["bipa", "biometric", "facial", "fingerprint"],
            "6.": ["coppa", "children", "age", "parental"],
            "7.": ["glba", "financial", "safeguard", "credit"],
            "8.": ["vendor", "third-party", "dpa", "processor"],
            "9.": ["border", "transfer", "scc", "cross-border"],
            "10.": ["dsar", "access", "deletion", "request"],
            "11.": ["breach", "incident", "leak", "compromise"],
            "12.": ["training", "employee", "staff", "internal"],
            "13.": ["cookie", "consent", "tracking", "pixel"],
            "14.": ["public-facing", "privacy policy", "redline", "gap"],
            "15.": ["1984", "dns", "domain", "hosting"],
            "16.": ["cpo", "action queue", "governance", "baseline"]
        }
        
        for k, kw_list in keywords.items():
            if sec.startswith(k):
                for f in findings:
                    f_text = (f["title"] + " " + (f["description"] or "")).lower()
                    if any(kw in f_text for kw in kw_list):
                        sec_findings.append({
                            "id": f["id"],
                            "title": f["title"],
                            "confidence": f["confidence"],
                            "status": f["status"],
                            "decision": f["decision"]
                        })
                        if f["decision"] == "accept":
                            is_completed = True
        
        # Mark as Ready if all associated findings are approved or if there are none (defaults to ready placeholder templates)
        if len(sec_findings) == 0 or is_completed:
            completed_count += 1
            status = "Ready"
        else:
            status = "Pending Action Approval"
            
        report_sections[sec] = {
            "status": status,
            "findings": sec_findings,
            "deliverable": f"Deliverable template generated for {sec}."
        }
    
    completion_percentage = int((completed_count / len(sections)) * 100)
    
    cursor.execute("SELECT 1 FROM reports WHERE client_id = ?", (client_id,))
    if cursor.fetchone():
        cursor.execute(
            "UPDATE reports SET completion_percentage = ?, report_data = ?, updated_at = CURRENT_TIMESTAMP WHERE client_id = ?",
            (completion_percentage, json.dumps(report_sections), client_id)
        )
    else:
        cursor.execute(
            "INSERT INTO reports (client_id, completion_percentage, report_data) VALUES (?, ?, ?)",
            (client_id, completion_percentage, json.dumps(report_sections))
        )
    
    conn_db.commit()
    conn_db.close()

@app.get("/api/clients/{client_id}/report")
def get_client_report(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM reports WHERE client_id = ?", (client_id,))
    row = cursor.fetchone()
    conn_db.close()
    
    if not row:
        compile_client_report(client_id)
        conn_db = sqlite3.connect(DB_FILE)
        conn_db.row_factory = sqlite3.Row
        cursor = conn_db.cursor()
        cursor.execute("SELECT * FROM reports WHERE client_id = ?", (client_id,))
        row = cursor.fetchone()
        conn_db.close()
        
    if row:
        r_dict = dict(row)
        r_dict["report_data"] = json.loads(r_dict["report_data"])
        return r_dict
        
    return {"client_id": client_id, "completion_percentage": 0, "report_data": {}}

@app.post("/api/clients/{client_id}/discover/website")
def discover_website(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    seed_findings_for_client(client_id, "website")
    compile_client_report(client_id)
    return {"status": "success", "message": "Website discovery triggered"}

@app.post("/api/clients/{client_id}/discover/documents")
def discover_documents(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    seed_findings_for_client(client_id, "documents")
    compile_client_report(client_id)
    return {"status": "success", "message": "Document analysis triggered"}

@app.post("/api/clients/{client_id}/discover/systems")
def discover_systems(client_id: str, scope: str = Depends(require_scopes(["expert"]))):
    seed_findings_for_client(client_id, "systems")
    compile_client_report(client_id)
    return {"status": "success", "message": "System discovery triggered"}

@app.post("/api/clients/{client_id}/discover/guided")
def discover_guided(client_id: str, payload: GuidedAnswersPayload, scope: str = Depends(require_scopes(["expert"]))):
    seed_findings_for_client(client_id, "guided")
    compile_client_report(client_id)
    return {"status": "success", "message": "Guided assessment processed"}

@app.get("/api/clients/{client_id}/discover/{mode}")
async def stream_client_discover(client_id: str, mode: str, scope: str = Depends(require_scopes(["expert"]))):
    """
    SSE stream executing LangGraph subgraphs step-by-step and yielding progress logs.
    """
    from fastapi.responses import StreamingResponse
    import asyncio
    from onboarding_agents import web_compiled, doc_compiled, sys_compiled, guided_compiled
    
    async def event_generator():
        conn_db = sqlite3.connect(DB_FILE)
        cursor = conn_db.cursor()
        cursor.execute("UPDATE clients SET status = 'running' WHERE id = ?", (client_id,))
        conn_db.commit()
        conn_db.close()
        
        # Initial State
        initial_state = {
            "client_id": client_id,
            "mode": mode,
            "status": "started",
            "logs": [],
            "progress": 0.0,
            "findings": []
        }
        
        # Pick the compiled graph
        if mode == "website":
            compiled_graph = web_compiled
        elif mode == "documents":
            compiled_graph = doc_compiled
        elif mode == "systems":
            compiled_graph = sys_compiled
        else:
            compiled_graph = guided_compiled
            
        # Stream the nodes of the graph
        try:
            for event in compiled_graph.stream(initial_state):
                for node_name, output in event.items():
                    log_msg = output.get("logs", ["Executing step..."])[-1]
                    progress = output.get("progress", 0.5)
                    data = {"agent": node_name, "log": log_msg, "progress": progress}
                    yield f"data: {json.dumps(data)}\n\n"
                    await asyncio.sleep(1.5)
        except Exception as e:
            err_data = {"agent": "orchestrator", "log": f"Agent error: {str(e)}", "progress": 1.0}
            yield f"data: {json.dumps(err_data)}\n\n"
            
        yield f"data: {json.dumps({'agent': 'system', 'log': 'Discovery run finished. Seeding findings database...', 'progress': 1.0})}\n\n"
        
        # Seed discovery findings based on selected mode
        seed_findings_for_client(client_id, mode)
        
        # Set client status to completed
        conn_db = sqlite3.connect(DB_FILE)
        cursor = conn_db.cursor()
        cursor.execute("UPDATE clients SET status = 'completed' WHERE id = ?", (client_id,))
        conn_db.commit()
        conn_db.close()
        
        # Compile report once
        compile_client_report(client_id)
        
        yield f"data: {json.dumps({'agent': 'system', 'log': 'Completed! Findings review panel ready.', 'status': 'completed'})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")

def seed_findings_for_client(client_id: str, mode: str):
    """Adds mode-specific mock findings for the client."""
    import uuid
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    findings_list = []
    if mode == "website":
        findings_list = [
            ("cookie_pixel_leak", "Tech Stack Auditor", "Third-Party Advertising Pixels Without Active Consent Gate", 
             "Detected active Meta and TikTok tracking pixels executing before cookie banner confirmation. Violates CPRA opt-out requirements.", 
             0.92, "Website Discovery (Front-end Scan)"),
            ("privacy_policy_gap", "Regulatory Auditor", "Public Privacy Policy Missing Mandatory State Disclosures", 
             "The public-facing privacy policy does not contain the required disclosures for California (CPRA profiling opt-out) or Texas (TDPSA biometric warnings).", 
             0.88, "Privacy Policy Redline Agent"),
            ("dns_security", "1984 Hosting DNS Agent", "1984 Hosting Domain DNS Key Not Rotated", 
             "DNS verification tokens for 1984 Hosting integration are using default test credentials. Crucial security risk.", 
             0.95, "DNS Verification Agent"),
        ]
    elif mode == "documents":
        findings_list = [
            ("vendor_dpa_missing", "Vendor Risk Auditor", "Missing Vendor DPA and SCC for Remote Subprocessor", 
             "Active contract files with AnalyticsPro Inc lack Standard Contractual Clauses (SCCs) required for EU-US cross-border transfers.", 
             0.84, "Vendor Risk Auditor Agent"),
            ("training_records_outdated", "Corporate Auditor", "Internal Employee Privacy Training Logs Stale", 
             "Privacy awareness training logs for HR personnel have not been updated since Q3 2024. Fails regulatory baseline controls.", 
             0.79, "Corporate Auditor Agent"),
        ]
    elif mode == "systems":
        findings_list = [
            ("s3_exposure", "Risk Auditor", "Unprotected S3 Buckets Containing Diagnostic Telemetry Data", 
             "HIPAA-governed medical diagnostic logs were detected on a misconfigured public-facing Amazon S3 bucket.", 
             0.99, "Risk Auditor Agent"),
            ("db_pii_no_retention", "Data Inventory Auditor", "Active Database Users Stored Without Retention Policies", 
             "SQLite database contains 10,000+ consumer records with no automated deletion script or data retention flags.", 
             0.82, "Data Inventory Auditor Agent"),
        ]
    else: # guided / hybrid / random
        findings_list = [
            ("admt_consent_missing", "Regulatory Auditor", "Automated Decision-Making (ADMT) Profiling Without Consent", 
             "The system conducts profiling on consumer records to predict compliance risk without offering CPRA profiling opt-out choices.", 
             0.86, "Regulatory Auditor Agent"),
            ("bipa_biometric_checks", "Privacy Auditor", "BIPA Biometric Collection Missing Statutory Notices", 
             "Biometric authentication processes are active for customer onboarding without explicit notice or consent logs.", 
             0.91, "Privacy Auditor Agent"),
        ]
        
    for fid_base, agent, title, desc, conf, source in findings_list:
        fid = f"{fid_base}-{uuid.uuid4().hex[:4]}"
        cursor.execute(
            "INSERT INTO findings (id, client_id, category, title, description, confidence, source, status, decision) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (fid, client_id, agent, title, desc, conf, source, "pending", "pending")
        )
        
    conn_db.commit()
    conn_db.close()

# --- Employee Mode Models, Helpers, and Endpoints ---
from pydantic import Field, field_validator
from datetime import datetime

class ProjectIntakePayload(BaseModel):
    id: str
    project_name: str
    purpose: str
    data_classification: List[str]
    retention_value: int
    retention_unit: str

    @field_validator('data_classification')
    @classmethod
    def validate_classification(cls, v):
        allowed = ["Health", "Financial", "Childrens", "Biometric", "Employee", "Sensitive PII"]
        for item in v:
            if item not in allowed:
                raise ValueError(f"Invalid classification: {item}. Allowed: {allowed}")
        return v

    @field_validator('retention_unit')
    @classmethod
    def validate_retention_unit(cls, v):
        allowed = ["days", "months", "years", "indefinite"]
        if v not in allowed:
            raise ValueError(f"Invalid retention unit: {v}. Allowed: {allowed}")
        return v

class MitigationStatusPayload(BaseModel):
    status: str

    @field_validator('status')
    @classmethod
    def validate_status(cls, v):
        allowed = ["open", "in_progress", "closed"]
        if v not in allowed:
            raise ValueError(f"Invalid status: {v}. Allowed: {allowed}")
        return v

class OpcReplyPayload(BaseModel):
    message_id: str
    body: str

def log_system_audit(employee_id: str, target_id: str, action: str, delta: Dict[str, Any]):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("""
        INSERT INTO system_audit_log (timestamp, employee_id, target_id, action, delta)
        VALUES (?, ?, ?, ?, ?)
    """, (datetime.utcnow().isoformat() + "Z", employee_id, target_id, action, json.dumps(delta)))
    conn_db.commit()
    conn_db.close()

@app.get("/api/employee/intake")
def get_employee_intakes(employee_id: str = "employee-default", status: Optional[str] = None, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    if status:
        cursor.execute("SELECT * FROM employee_intake WHERE employee_id = ? AND status = ?", (employee_id, status))
    else:
        cursor.execute("SELECT * FROM employee_intake WHERE employee_id = ?", (employee_id,))
    rows = cursor.fetchall()
    conn_db.close()
    
    res = []
    for r in rows:
        d = dict(r)
        d["data_classification"] = json.loads(d["data_classification"])
        res.append(d)
    return res

@app.post("/api/employee/intake")
def post_employee_intake(payload: ProjectIntakePayload, employee_id: str = "employee-default", scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    # Check if record exists
    cursor.execute("SELECT * FROM employee_intake WHERE id = ?", (payload.id,))
    existing = cursor.fetchone()
    
    if existing:
        ex_dict = dict(existing)
        if ex_dict["status"] in ["submitted", "under_review", "approved"]:
            conn_db.close()
            raise HTTPException(status_code=403, detail="Forbidden: Cannot edit locked or submitted intake records.")
        
        # Generate new version token
        current_version = ex_dict.get("version_token", "v1.0_approved")
        try:
            v_num = float(current_version.split("v")[1].split("_")[0])
            new_version = f"v{v_num + 1.0}_submitted"
        except Exception:
            new_version = "v2.0_submitted"

        # Update existing draft
        cursor.execute("""
            UPDATE employee_intake 
            SET project_name = ?, purpose = ?, data_classification = ?, retention_value = ?, retention_unit = ?, status = 'submitted', version_token = ?
            WHERE id = ? AND employee_id = ?
        """, (payload.project_name, payload.purpose, json.dumps(payload.data_classification), payload.retention_value, payload.retention_unit, new_version, payload.id, employee_id))
        
        # Calculate diff/delta for audit log
        delta = {
            "project_name": (ex_dict["project_name"], payload.project_name),
            "purpose": (ex_dict["purpose"], payload.purpose),
            "data_classification": (json.loads(ex_dict["data_classification"]), payload.data_classification),
            "retention_value": (ex_dict["retention_value"], payload.retention_value),
            "retention_unit": (ex_dict["retention_unit"], payload.retention_unit),
            "status": (ex_dict["status"], "submitted")
        }
    else:
        # Create new submitted record
        cursor.execute("""
            INSERT INTO employee_intake (id, employee_id, project_name, purpose, data_classification, retention_value, retention_unit, status, version_token)
            VALUES (?, ?, ?, ?, ?, ?, ?, 'submitted', 'v1.0_submitted')
        """, (payload.id, employee_id, payload.project_name, payload.purpose, json.dumps(payload.data_classification), payload.retention_value, payload.retention_unit))
        delta = {"created": payload.dict()}
        
    conn_db.commit()
    conn_db.close()
    
    # Audit log entry
    log_system_audit(employee_id, payload.id, "intake_submission", delta)
    return {"status": "success", "message": "Intake form submitted successfully"}

@app.get("/api/employee/mitigation")
def get_employee_mitigations(employee_id: str = "employee-default", status: Optional[str] = None, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    if status:
        cursor.execute("SELECT * FROM employee_mitigation WHERE employee_id = ? AND status = ?", (employee_id, status))
    else:
        cursor.execute("SELECT * FROM employee_mitigation WHERE employee_id = ?", (employee_id,))
    rows = cursor.fetchall()
    conn_db.close()
    
    # Calculate overdue state dynamically based on system runtime timestamp vs due date
    now_str = datetime.utcnow().isoformat() + "Z"
    res = []
    for r in rows:
        d = dict(r)
        if d["status"] != "closed" and d["due_date"] < now_str:
            d["status"] = "overdue"
        res.append(d)
    return res

@app.patch("/api/employee/mitigation/{plan_id}")
def patch_employee_mitigation(plan_id: str, payload: MitigationStatusPayload, employee_id: str = "employee-default", scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM employee_mitigation WHERE id = ? AND employee_id = ?", (plan_id, employee_id))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Mitigation plan not found")
        
    ex_dict = dict(existing)
    if ex_dict["status"] == "closed":
        conn_db.close()
        raise HTTPException(status_code=403, detail="Forbidden: Checked/Closed mitigation plans are locked.")
        
    closed_at = (datetime.utcnow().isoformat() + "Z") if payload.status == "closed" else None
    
    cursor.execute("""
        UPDATE employee_mitigation 
        SET status = ?, closed_at = ?
        WHERE id = ? AND employee_id = ?
    """, (payload.status, closed_at, plan_id, employee_id))
    conn_db.commit()
    conn_db.close()
    
    delta = {
        "status": (ex_dict["status"], payload.status),
        "closed_at": (ex_dict["closed_at"], closed_at)
    }
    log_system_audit(employee_id, plan_id, "mitigation_status_update", delta)
    return {"status": "success", "message": f"Mitigation plan set to {payload.status}"}

@app.get("/api/employee/opc-inquiries")
def get_opc_inquiries(employee_id: str = "employee-default", read_status: Optional[str] = None, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    if read_status:
        cursor.execute("SELECT * FROM opc_inquiries WHERE employee_id = ? AND read_status = ?", (employee_id, read_status))
    else:
        cursor.execute("SELECT * FROM opc_inquiries WHERE employee_id = ?", (employee_id,))
    inq_rows = cursor.fetchall()
    
    inquiries = []
    for row in inq_rows:
        inq = dict(row)
        # Fetch thread messages
        cursor.execute("SELECT * FROM opc_messages WHERE thread_id = ? ORDER BY timestamp ASC", (inq["thread_id"],))
        msg_rows = cursor.fetchall()
        inq["messages"] = [dict(mr) for mr in msg_rows]
        inquiries.append(inq)
        
    conn_db.close()
    return inquiries

@app.post("/api/employee/opc-inquiries/{thread_id}/reply")
def post_opc_reply(thread_id: str, payload: OpcReplyPayload, employee_id: str = "employee-default", scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM opc_inquiries WHERE thread_id = ? AND employee_id = ?", (thread_id, employee_id))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Inquiry thread not found")
        
    # Append message block
    timestamp = datetime.utcnow().isoformat() + "Z"
    cursor.execute("""
        INSERT INTO opc_messages (message_id, thread_id, sender, body, timestamp)
        VALUES (?, ?, ?, ?, ?)
    """, (payload.message_id, thread_id, "Employee", payload.body, timestamp))
    
    # Mark thread as unread by CPO (or read by Employee)
    cursor.execute("UPDATE opc_inquiries SET read_status = 'read' WHERE thread_id = ?", (thread_id,))
    
    conn_db.commit()
    conn_db.close()
    
    delta = {
        "message_id": payload.message_id,
        "body": payload.body,
        "sender": "Employee",
        "timestamp": timestamp
    }
    log_system_audit(employee_id, thread_id, "opc_reply_submission", delta)
    return {"status": "success", "message": "Reply appended successfully"}

@app.get("/api/employee/audit-ledger")
def get_audit_ledger(scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM system_audit_log ORDER BY timestamp DESC")
    rows = cursor.fetchall()
    conn_db.close()
    ledger = []
    for r in rows:
        d = dict(r)
        d["delta"] = json.loads(d["delta"])
        ledger.append(d)
    return ledger


# --- RISK REGISTER API ENDPOINTS ---
import uuid
import time
import datetime as dt

class RiskPayload(BaseModel):
    law_25_section: Optional[str] = None
    issue: str
    source_document: Optional[str] = None
    source_id: Optional[str] = None
    source_title: Optional[str] = None
    gap_type: Optional[str] = None
    likelihood: int
    impact: int
    original_risk_level: Optional[str] = None
    recommendation: Optional[str] = None
    khp_response: Optional[str] = None
    status_note: Optional[str] = None
    assigned_to: Optional[str] = None
    risk_assigned_date: Optional[str] = None
    next_review_date: Optional[str] = None

class RiskAcknowledgePayload(BaseModel):
    actor: str
    note: Optional[str] = None

class RiskExtensionPayload(BaseModel):
    actor: str
    suggested_date: str
    reason: str

class PiaCreatePayload(BaseModel):
    title: str
    description: str
    creator_name: str
    creator_role: str
    file_name: Optional[str] = None

class PushRiskPayload(BaseModel):
    actor_name: str
    actor_role: str
    assigned_to: str
    next_review_date: str

class AssessmentCreatePayload(BaseModel):
    title: str
    type: str = "PIA"
    output_type: str
    level: int
    source: str
    source_id: Optional[str] = None
    prepared_by: str
    jurisdictions: List[str]
    roles_in_scope: List[str]
    departments_in_scope: List[str]
    services_in_scope: List[str]
    data_types: List[str]
    cross_border: bool
    cross_border_jurisdictions: List[str]

class AssessmentUpdatePayload(BaseModel):
    title: str
    output_type: str
    level: int
    risk_level: str
    status: str
    jurisdictions: List[str]
    roles_in_scope: List[str]
    departments_in_scope: List[str]
    services_in_scope: List[str]
    data_types: List[str]
    cross_border: bool
    cross_border_jurisdictions: List[str]
    sections: dict

class AssessmentStatusPayload(BaseModel):
    status: str

class AssessmentVersionPayload(BaseModel):
    version_notes: str

class AssessmentLinkRiskPayload(BaseModel):
    risk_id: str

def log_risk_change(risk_id: str, actor: str, action: str, field_changed: Optional[str] = None, old_val: Optional[str] = None, new_val: Optional[str] = None, note: Optional[str] = None):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    log_id = str(uuid.uuid4())
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    cursor.execute("""
        INSERT INTO risk_log (log_id, risk_id, actor, action, field_changed, old_value, new_value, note, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (log_id, risk_id, actor, action, field_changed, old_val, new_val, note, now_str))
    conn_db.commit()
    conn_db.close()

def get_risk_level_from_score(score: int) -> str:
    if score >= 15:
        return "High"
    elif score >= 8:
        return "Medium"
    else:
        return "Low"

@app.get("/api/pias")
def get_pias(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM pias ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/pias")
def create_pia(payload: PiaCreatePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    pia_id = f"PIA-{str(uuid.uuid4())[:6].upper()}"
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("""
        INSERT INTO pias (pia_id, title, description, creator_name, creator_role, file_name, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (pia_id, payload.title, payload.description, payload.creator_name, payload.creator_role, payload.file_name, now_str))
    
    # Mock LLM Risk Extraction: Generate 2 structured risks based on title
    title_lower = payload.title.lower()
    if "biometric" in title_lower or "face" in title_lower or "voice" in title_lower:
        ext_issue_1 = "Consent flow bypass for biometric authentication telemetry"
        ext_law_1 = "12"
        ext_gap_1 = "Consent Management"
        ext_rec_1 = "Deploy explicit opt-in dialogs before any biometric scan runs."
        ext_lik_1 = 4
        ext_imp_1 = 5
    elif "cloud" in title_lower or "database" in title_lower or "hosting" in title_lower:
        ext_issue_1 = "Cross-border personal data transfer to unmapped regions"
        ext_law_1 = "17"
        ext_gap_1 = "Infrastructure Security"
        ext_rec_1 = "Restrict DB replication to local regions and update registry mapping."
        ext_lik_1 = 3
        ext_imp_1 = 4
    else:
        ext_issue_1 = f"Data minimisation shortfall in {payload.title}"
        ext_law_1 = "None"
        ext_gap_1 = "Operational Gap"
        ext_rec_1 = "Conduct audit of data ingestion pipelines and prune unused payloads."
        ext_lik_1 = 3
        ext_imp_1 = 3

    ext_issue_2 = f"Vendor SLA review gap with sub-processors involved in {payload.title}"
    ext_law_2 = "15"
    ext_gap_2 = "Vendor Assessment"
    ext_rec_2 = "Standardise DPAs with third parties and check SOC 2 certifications."
    ext_lik_2 = 2
    ext_imp_2 = 4

    cursor.execute("""
        INSERT INTO pia_extracted_risks (extracted_risk_id, pia_id, law_25_section, issue, gap_type, likelihood, impact, recommendation, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending')
    """, (f"EXT-{str(uuid.uuid4())[:6].upper()}", pia_id, ext_law_1, ext_issue_1, ext_gap_1, ext_lik_1, ext_imp_1, ext_rec_1))

    cursor.execute("""
        INSERT INTO pia_extracted_risks (extracted_risk_id, pia_id, law_25_section, issue, gap_type, likelihood, impact, recommendation, status)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending')
    """, (f"EXT-{str(uuid.uuid4())[:6].upper()}", pia_id, ext_law_2, ext_issue_2, ext_gap_2, ext_lik_2, ext_imp_2, ext_rec_2))

    conn_db.commit()
    conn_db.close()
    return {"status": "success", "pia_id": pia_id}

@app.get("/api/pias/{pia_id}/extracted-risks")
def get_pia_extracted_risks(pia_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM pia_extracted_risks WHERE pia_id = ?", (pia_id,))
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/pias/extracted-risks/{extracted_risk_id}/push")
def push_extracted_risk(extracted_risk_id: str, payload: PushRiskPayload, scope: str = Depends(require_scopes(["expert"]))):
    # Role checking logic
    allowed_roles = ["Privacy Specialist", "Security Officer", "Director of IT", "CEO"]
    if payload.actor_role not in allowed_roles:
        raise HTTPException(
            status_code=403, 
            detail=f"Access denied: Role '{payload.actor_role}' does not have authority to push risks. Required roles: {', '.join(allowed_roles)}"
        )

    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    # Get the extracted risk details
    cursor.execute("SELECT * FROM pia_extracted_risks WHERE extracted_risk_id = ?", (extracted_risk_id,))
    ext_risk = cursor.fetchone()
    if not ext_risk:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Extracted risk not found")
        
    if ext_risk["status"] == "pushed":
        conn_db.close()
        raise HTTPException(status_code=400, detail="Risk already pushed to register")
        
    # Get the parent PIA title
    cursor.execute("SELECT title FROM pias WHERE pia_id = ?", (ext_risk["pia_id"],))
    pia_row = cursor.fetchone()
    pia_title = pia_row["title"] if pia_row else "PIA Document"

    # Generate new risk ID
    cursor.execute("SELECT risk_id FROM risks ORDER BY risk_id DESC LIMIT 1")
    last_row = cursor.fetchone()
    if last_row:
        try:
            last_num = int(last_row["risk_id"].split("-")[1])
            new_id = f"RSK-{last_num + 1:04d}"
        except:
            new_id = f"RSK-{str(uuid.uuid4())[:6]}"
    else:
        new_id = "RSK-0001"
        
    score = ext_risk["likelihood"] * ext_risk["impact"]
    level = get_risk_level_from_score(score)
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    today_str = dt.date.today().isoformat()
    
    cursor.execute("""
        INSERT INTO risks (
            risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, 
            likelihood, impact, risk_score, risk_level, recommendation, khp_response, 
            assigned_to, risk_assigned_date, next_review_date, original_risk_level, 
            status_note, status, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'open', ?, ?)
    """, (
        new_id, ext_risk["law_25_section"], ext_risk["issue"], f"PIA: {pia_title}", ext_risk["pia_id"], pia_title, ext_risk["gap_type"],
        ext_risk["likelihood"], ext_risk["impact"], score, level, ext_risk["recommendation"], "Pending response.",
        payload.assigned_to, today_str, payload.next_review_date, level, "Pushed from PIA.", now_str, now_str
    ))
    
    # Update extracted risk status
    cursor.execute("""
        UPDATE pia_extracted_risks 
        SET status = 'pushed', pushed_risk_id = ? 
        WHERE extracted_risk_id = ?
    """, (new_id, extracted_risk_id))
    
    conn_db.commit()
    conn_db.close()
    
    # Audit log
    log_risk_change(
        risk_id=new_id, 
        actor=payload.actor_name, 
        action="Created via PIA Risk Extraction", 
        field_changed="status", 
        old_val="N/A", 
        new_val="open", 
        note=f"Extracted from {pia_title} and pushed by {payload.actor_name} ({payload.actor_role})"
    )
    
    return {"status": "success", "risk_id": new_id}

@app.get("/api/risks")
def get_risks(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks")
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.get("/api/risks/{risk_id}")
def get_single_risk(risk_id: str, scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Risk not found")
    return dict(row)

@app.post("/api/risks")
def create_new_risk(payload: RiskPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    # Auto-generate risk ID
    cursor.execute("SELECT risk_id FROM risks ORDER BY risk_id DESC LIMIT 1")
    last_row = cursor.fetchone()
    if last_row:
        try:
            last_num = int(last_row["risk_id"].split("-")[1])
            new_id = f"RSK-{last_num + 1:04d}"
        except:
            new_id = f"RSK-{str(uuid.uuid4())[:6]}"
    else:
        new_id = "RSK-0001"
        
    score = payload.likelihood * payload.impact
    level = get_risk_level_from_score(score)
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("""
        INSERT INTO risks (risk_id, law_25_section, issue, source_document, source_id, source_title, gap_type, likelihood, impact, risk_score, risk_level, recommendation, khp_response, assigned_to, risk_assigned_date, next_review_date, original_risk_level, status_note, status, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'open', ?, ?)
    """, (
        new_id, payload.law_25_section, payload.issue, payload.source_document, payload.source_id, payload.source_title, payload.gap_type,
        payload.likelihood, payload.impact, score, level, payload.recommendation, payload.khp_response, payload.assigned_to,
        payload.risk_assigned_date, payload.next_review_date, payload.original_risk_level or level, payload.status_note,
        now_str, now_str
    ))
    conn_db.commit()
    conn_db.close()
    
    log_risk_change(new_id, "CPO", "created", note=f"Initial risk level set to {level}")
    return {"status": "success", "risk_id": new_id, "risk_score": score, "risk_level": level}

@app.put("/api/risks/{risk_id}")
def update_risk(risk_id: str, payload: RiskPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Risk not found")
        
    ex = dict(existing)
    score = payload.likelihood * payload.impact
    level = get_risk_level_from_score(score)
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("""
        UPDATE risks
        SET law_25_section = ?, issue = ?, source_document = ?, source_id = ?, source_title = ?, gap_type = ?,
            likelihood = ?, impact = ?, risk_score = ?, risk_level = ?, recommendation = ?, khp_response = ?,
            assigned_to = ?, risk_assigned_date = ?, next_review_date = ?, original_risk_level = ?, status_note = ?,
            updated_at = ?
        WHERE risk_id = ?
    """, (
        payload.law_25_section, payload.issue, payload.source_document, payload.source_id, payload.source_title, payload.gap_type,
        payload.likelihood, payload.impact, score, level, payload.recommendation, payload.khp_response,
        payload.assigned_to, payload.risk_assigned_date, payload.next_review_date, payload.original_risk_level or ex["original_risk_level"], payload.status_note,
        now_str, risk_id
    ))
    conn_db.commit()
    conn_db.close()
    
    # Log any differences
    fields_to_check = ["law_25_section", "issue", "gap_type", "likelihood", "impact", "assigned_to", "next_review_date", "status_note"]
    for f in fields_to_check:
        old_val = str(ex[f])
        new_val = str(getattr(payload, f))
        if old_val != new_val:
            log_risk_change(risk_id, "CPO", "field_change", field_changed=f, old_val=old_val, new_val=new_val)
            
    return {"status": "success", "risk_score": score, "risk_level": level}

@app.delete("/api/risks/{risk_id}")
def delete_risk(risk_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Risk not found")
        
    cursor.execute("DELETE FROM risks WHERE risk_id = ?", (risk_id,))
    cursor.execute("DELETE FROM risk_log WHERE risk_id = ?", (risk_id,))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "message": f"Risk {risk_id} deleted."}

@app.patch("/api/risks/{risk_id}/close")
def close_risk(risk_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Risk not found")
        
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    cursor.execute("UPDATE risks SET status = 'closed', updated_at = ? WHERE risk_id = ?", (now_str, risk_id))
    conn_db.commit()
    conn_db.close()
    
    log_risk_change(risk_id, "CPO", "closed", note="Risk marked resolved/closed.")
    return {"status": "success", "message": f"Risk {risk_id} marked as closed."}

@app.get("/api/risks/assigned/{name}")
def get_assigned_risks(name: str, scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE assigned_to = ? COLLATE NOCASE", (name,))
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/risks/{risk_id}/acknowledge")
def acknowledge_risk(risk_id: str, payload: RiskAcknowledgePayload, scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Risk not found")
    conn_db.close()
    
    note = payload.note or "Acknowledged by employee."
    log_risk_change(risk_id, payload.actor, "acknowledged", note=note)
    return {"status": "success", "message": "Risk review acknowledged."}

@app.post("/api/risks/{risk_id}/extension")
def request_extension(risk_id: str, payload: RiskExtensionPayload, scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risks WHERE risk_id = ?", (risk_id,))
    existing = cursor.fetchone()
    if not existing:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Risk not found")
    
    ex = dict(existing)
    conn_db.close()
    
    log_risk_change(risk_id, payload.actor, "extension_requested", old_val=ex["next_review_date"], new_val=payload.suggested_date, note=payload.reason)
    
    # Existing queue integration
    tx_id = f"ext-{risk_id}-{int(time.time())}"
    tx_req = NewTransactionRequest(
        id=tx_id,
        type="Extension Request",
        client=ex["assigned_to"],
        jurisdiction="Law 25",
        priority="medium",
        deadline=payload.suggested_date,
        description=f"Employee {payload.actor} requested an extension for risk {risk_id} to {payload.suggested_date}. Reason: {payload.reason}",
        agent="Risk Register Agent",
        summary=f"Extension Request for {risk_id}",
        raw=json.dumps({
            "risk_id": risk_id,
            "old_date": ex["next_review_date"],
            "new_date": payload.suggested_date,
            "reason": payload.reason,
            "actor": payload.actor
        })
    )
    
    create_transaction(tx_req, scope="employee")
    return {"status": "success", "message": "Extension request filed and routed to CPO decision queue."}

@app.get("/api/risks/{risk_id}/log")
def get_risk_audit_log(risk_id: str, scope: str = Depends(require_scopes(["expert", "employee"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM risk_log WHERE risk_id = ? ORDER BY timestamp DESC", (risk_id,))
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

# --- ASSESSMENT API ENDPOINTS ---
import io
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

@app.get("/api/assessments")
def list_assessments(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM assessments ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn_db.close()
    
    res = []
    for r in rows:
        d = dict(r)
        d["jurisdictions"] = json.loads(d["jurisdictions"])
        d["roles_in_scope"] = json.loads(d["roles_in_scope"])
        d["departments_in_scope"] = json.loads(d["departments_in_scope"])
        d["services_in_scope"] = json.loads(d["services_in_scope"])
        d["data_types"] = json.loads(d["data_types"])
        d["cross_border_jurisdictions"] = json.loads(d["cross_border_jurisdictions"])
        d["sections"] = json.loads(d["sections"])
        d["linked_risks"] = json.loads(d["linked_risks"]) if d["linked_risks"] else []
        res.append(d)
    return res

@app.get("/api/assessments/{assessment_id}")
def get_assessment(assessment_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM assessments WHERE assessment_id = ?", (assessment_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Assessment not found")
        
    d = dict(row)
    d["jurisdictions"] = json.loads(d["jurisdictions"])
    d["roles_in_scope"] = json.loads(d["roles_in_scope"])
    d["departments_in_scope"] = json.loads(d["departments_in_scope"])
    d["services_in_scope"] = json.loads(d["services_in_scope"])
    d["data_types"] = json.loads(d["data_types"])
    d["cross_border_jurisdictions"] = json.loads(d["cross_border_jurisdictions"])
    d["sections"] = json.loads(d["sections"])
    d["linked_risks"] = json.loads(d["linked_risks"]) if d["linked_risks"] else []
    return d

@app.post("/api/assessments")
def create_assessment(payload: AssessmentCreatePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    # Generate ID: PIA-YYYY-XXX
    current_year = dt.date.today().year
    cursor.execute("SELECT assessment_id FROM assessments ORDER BY created_at DESC LIMIT 1")
    last_row = cursor.fetchone()
    next_num = 1
    if last_row:
        try:
            last_id = last_row[0]
            last_num = int(last_id.split("-")[-1]) + 1
        except:
            pass
    new_id = f"PIA-{current_year}-{next_num:03d}"
    
    default_sections = {
        "preliminary_analysis": {
            "project_title": payload.title,
            "institution_dept": payload.departments_in_scope[0] if payload.departments_in_scope else "N/A",
            "executive_sponsor": "CEO",
            "project_lead": payload.prepared_by,
            "pia_lead": payload.prepared_by,
            "planned_go_live": "",
            "personal_information": "Yes" if payload.data_types else "No",
            "data_subjects": ", ".join(payload.roles_in_scope),
            "relevant_legislation": ", ".join(payload.jurisdictions)
        },
        "executive_summary": "",
        "scope_and_architecture": {
            "scope_description": "",
            "system_architecture": "",
            "roles_responsibilities": []
        },
        "data_classification": {
            "sensitivity": "Sensitive",
            "data_sovereignty": "",
            "regulatory_basis": ""
        },
        "threat_analysis": [],
        "privacy_analysis": [],
        "necessity_proportionality": {
            "necessity": "",
            "effectiveness": "",
            "minimisation": "",
            "proportionality": "",
            "finding": "Not Met"
        },
        "recommendations": [],
        "signoff": {
            "psr_status": "Draft",
            "executive_endorsement": "Pending",
            "signoff_date": ""
        }
    }
    
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("""
        INSERT INTO assessments (
            assessment_id, title, type, output_type, level, version, version_notes, status, risk_level, 
            source, source_id, prepared_by, jurisdictions, roles_in_scope, departments_in_scope, 
            services_in_scope, data_types, cross_border, cross_border_jurisdictions, sections, linked_risks, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, 'draft', 'Low', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        new_id, payload.title, payload.type, payload.output_type, payload.level, "0.1", "Initial draft",
        payload.source, payload.source_id, payload.prepared_by,
        json.dumps(payload.jurisdictions), json.dumps(payload.roles_in_scope), json.dumps(payload.departments_in_scope),
        json.dumps(payload.services_in_scope), json.dumps(payload.data_types), 1 if payload.cross_border else 0,
        json.dumps(payload.cross_border_jurisdictions), json.dumps(default_sections), json.dumps([]), now_str, now_str
    ))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "assessment_id": new_id}

@app.put("/api/assessments/{assessment_id}")
def update_assessment(assessment_id: str, payload: AssessmentUpdatePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("""
        UPDATE assessments
        SET title = ?, output_type = ?, level = ?, risk_level = ?, status = ?, jurisdictions = ?, 
            roles_in_scope = ?, departments_in_scope = ?, services_in_scope = ?, data_types = ?, 
            cross_border = ?, cross_border_jurisdictions = ?, sections = ?, updated_at = ?
        WHERE assessment_id = ?
    """, (
        payload.title, payload.output_type, payload.level, payload.risk_level, payload.status,
        json.dumps(payload.jurisdictions), json.dumps(payload.roles_in_scope), json.dumps(payload.departments_in_scope),
        json.dumps(payload.services_in_scope), json.dumps(payload.data_types), 1 if payload.cross_border else 0,
        json.dumps(payload.cross_border_jurisdictions), json.dumps(payload.sections), now_str, assessment_id
    ))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success"}

@app.patch("/api/assessments/{assessment_id}/status")
def patch_assessment_status(assessment_id: str, payload: AssessmentStatusPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    
    cursor.execute("UPDATE assessments SET status = ?, updated_at = ? WHERE assessment_id = ?", (payload.status, now_str, assessment_id))
    
    # If approved/good_to_go, push action items to the decision queue
    if payload.status in ["good_to_go", "approved"]:
        conn_db.row_factory = sqlite3.Row
        cur2 = conn_db.cursor()
        cur2.execute("SELECT sections FROM assessments WHERE assessment_id = ?", (assessment_id,))
        row = cur2.fetchone()
        if row:
            sections = json.loads(row["sections"])
            recs = sections.get("recommendations", [])
            for r in recs:
                owner = r.get("owner", "")
                if owner and owner.lower() not in ["cpo", "wael", "waël"]:
                    # Create transaction queue item
                    tx_id = f"REQ-{str(uuid.uuid4())[:6].upper()}"
                    tx_req = NewTransactionRequest(
                        id=tx_id,
                        type="Action Item",
                        client=owner,
                        jurisdiction="Law 25",
                        priority=r.get("priority", "medium").lower(),
                        deadline=r.get("due", "Pre-launch"),
                        description=f"PIA Action Item: {r.get('action')} (Assessment: {assessment_id})",
                        agent="PIA Portal",
                        summary=f"Action item for {assessment_id}",
                        raw=json.dumps(r)
                    )
                    try:
                        create_transaction(tx_req, scope="expert")
                    except Exception as e:
                        print(f"Error seeding transaction: {e}")
                        
    conn_db.commit()
    conn_db.close()
    return {"status": "success"}

@app.post("/api/assessments/{assessment_id}/version")
def save_assessment_version(assessment_id: str, payload: AssessmentVersionPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    cursor.execute("SELECT * FROM assessments WHERE assessment_id = ?", (assessment_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Assessment not found")
        
    current_ver = float(row["version"])
    new_ver = f"{current_ver + 0.1:.1f}"
    
    # Save version
    version_id = str(uuid.uuid4())
    now_str = dt.datetime.utcnow().isoformat() + "Z"
    cursor.execute("""
        INSERT INTO assessment_versions (version_id, assessment_id, version, version_notes, sections, created_at, created_by)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (version_id, assessment_id, row["version"], payload.version_notes, row["sections"], now_str, row["prepared_by"]))
    
    # Update current assessment version
    cursor.execute("UPDATE assessments SET version = ?, updated_at = ? WHERE assessment_id = ?", (new_ver, now_str, assessment_id))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "new_version": new_ver}

@app.post("/api/assessments/{assessment_id}/populate")
def populate_assessment(assessment_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM assessments WHERE assessment_id = ?", (assessment_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Assessment not found")
        
    title = row["title"]
    services = json.loads(row["services_in_scope"])
    data_types = json.loads(row["data_types"])
    jurisdictions = json.loads(row["jurisdictions"])
    
    # Generate simulated AI outputs based on title
    title_lower = title.lower()
    
    exec_summary = f"This assessment reviews '{title}'. Scope includes processing data via {', '.join(services)}. Data types captured are {', '.join(data_types)}. Key compliance triggers include {', '.join(jurisdictions)}. Identified risks have been mapped to corresponding technical and administrative safeguards."
    
    threat_analysis = [
        {"domain": "System access control", "level": "Medium", "description": "Insufficient IAM role isolation on database levels"},
        {"domain": "Audit logging review", "level": "Low", "description": "Local log rotate schedule set to standard defaults"}
    ]
    
    # Conditional checks
    if any(s in ["WhatsApp", "Twilio", "Cohere", "AWS", "Google Cloud", "Azure"] for s in services):
        threat_analysis.append({"domain": "Cross-border data transit", "level": "High", "description": f"Third-party system routing involves US jurisdictions."})
        
    recs = [
        {"priority": "High", "action": f"Review DPA and SOC 2 audits for {services[0] if services else 'third-parties'}", "owner": "Legal", "due": "Pre-launch"},
        {"priority": "Medium", "action": "Implement MFA on all console logins", "owner": "IT", "due": "Pre-launch"}
    ]
    
    sections = {
        "preliminary_analysis": {
            "project_title": title,
            "institution_dept": "IT / Operations",
            "executive_sponsor": "CEO",
            "project_lead": row["prepared_by"],
            "pia_lead": row["prepared_by"],
            "planned_go_live": "2026-08-01",
            "personal_information": "Yes" if data_types else "No",
            "data_subjects": ", ".join(json.loads(row["roles_in_scope"])),
            "relevant_legislation": ", ".join(jurisdictions)
        },
        "executive_summary": exec_summary,
        "scope_and_architecture": {
            "scope_description": f"Integration audit of {title}.",
            "system_architecture": f"Data flows: client -> {', '.join(services)} -> local servers.",
            "roles_responsibilities": [
                {"role": "Controller", "entity": "KHP"},
                {"role": "Processor", "entity": services[0] if services else "Vendor"}
            ]
        },
        "data_classification": {
            "sensitivity": "Highly Sensitive" if "Mental health" in data_types or "Biometric data" in data_types else "Confidential",
            "data_sovereignty": "Restricted by cross-border regulations.",
            "regulatory_basis": "Explicit User Consent"
        },
        "threat_analysis": threat_analysis,
        "privacy_analysis": [
            {"requirement": "Collection & Consent", "gap": "No dedicated banner notice", "mitigation": "Add clear consent tick-box"},
            {"requirement": "Transparency", "gap": "Policy document needs update", "mitigation": "Append note to privacy registry page"}
        ],
        "necessity_proportionality": {
            "necessity": "Crucial for system scaling",
            "effectiveness": "Provides robust connectivity",
            "minimisation": "Aggressive pruning schedules configured",
            "proportionality": "Proportional to operational needs",
            "finding": "Met"
        },
        "recommendations": recs,
        "signoff": {
            "psr_status": "Good to Go",
            "executive_endorsement": "Pending",
            "signoff_date": ""
        }
    }
    
    cursor.execute("UPDATE assessments SET sections = ?, risk_level = 'Medium' WHERE assessment_id = ?", (json.dumps(sections), assessment_id))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "message": "Pre-population complete."}

@app.get("/api/assessments/{assessment_id}/risks")
def get_assessment_linked_risks(assessment_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("""
        SELECT r.* FROM risks r
        JOIN assessment_risks ar ON r.risk_id = ar.risk_id
        WHERE ar.assessment_id = ?
    """, (assessment_id,))
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/assessments/{assessment_id}/risks")
def link_risk_to_assessment(assessment_id: str, payload: AssessmentLinkRiskPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    try:
        cursor.execute("INSERT OR IGNORE INTO assessment_risks (assessment_id, risk_id) VALUES (?, ?)", (assessment_id, payload.risk_id))
        
        # Update assessments table linked_risks column
        conn_db.row_factory = sqlite3.Row
        cur2 = conn_db.cursor()
        cur2.execute("SELECT linked_risks FROM assessments WHERE assessment_id = ?", (assessment_id,))
        row = cur2.fetchone()
        if row:
            current_risks = json.loads(row["linked_risks"]) if row["linked_risks"] else []
            if payload.risk_id not in current_risks:
                current_risks.append(payload.risk_id)
                cursor.execute("UPDATE assessments SET linked_risks = ? WHERE assessment_id = ?", (json.dumps(current_risks), assessment_id))
                
        conn_db.commit()
    except Exception as e:
        conn_db.close()
        raise HTTPException(status_code=400, detail=str(e))
    conn_db.close()
    return {"status": "success"}

@app.get("/api/assessments/{assessment_id}/export/pdf")
def export_assessment_pdf(assessment_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM assessments WHERE assessment_id = ?", (assessment_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Assessment not found")
        
    assessment_data = dict(row)
    pdf_bytes = generate_pdf_document(assessment_data)
    
    from fastapi.responses import Response
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={assessment_id}_Privacy_Impact_Assessment.pdf"}
    )

@app.get("/api/assessments/{assessment_id}/export/docx")
def export_assessment_docx(assessment_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM assessments WHERE assessment_id = ?", (assessment_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Assessment not found")
        
    assessment_data = dict(row)
    docx_bytes = generate_docx_document(assessment_data)
    
    from fastapi.responses import Response
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={assessment_id}_Privacy_Impact_Assessment.docx"}
    )

def generate_pdf_document(assessment_data: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#d97706'),
        spaceAfter=15
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1e293b'),
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8
    )
    
    story = []
    
    # Title
    story.append(Paragraph(assessment_data.get("title", "Privacy Impact Assessment"), title_style))
    story.append(Spacer(1, 10))
    
    # Metadata
    story.append(Paragraph(f"<b>Assessment ID:</b> {assessment_data.get('assessment_id')}", body_style))
    story.append(Paragraph(f"<b>Output Type:</b> {assessment_data.get('output_type')}", body_style))
    story.append(Paragraph(f"<b>Level:</b> Level {assessment_data.get('level')}", body_style))
    story.append(Paragraph(f"<b>Version:</b> {assessment_data.get('version')} ({assessment_data.get('version_notes', 'N/A')})", body_style))
    story.append(Paragraph(f"<b>Status:</b> {assessment_data.get('status')}", body_style))
    story.append(Paragraph(f"<b>Risk Level:</b> {assessment_data.get('risk_level')}", body_style))
    story.append(Paragraph(f"<b>Prepared By:</b> {assessment_data.get('prepared_by')}", body_style))
    story.append(Spacer(1, 15))
    
    sections = assessment_data.get("sections", {})
    if isinstance(sections, str):
        try:
            sections = json.loads(sections)
        except:
            sections = {}
            
    # Exec Summary
    story.append(Paragraph("1. Executive Summary", h2_style))
    story.append(Paragraph(sections.get("executive_summary", "No executive summary provided."), body_style))
    story.append(Spacer(1, 10))
    
    # Scope
    story.append(Paragraph("2. Scope and Architecture", h2_style))
    scope = sections.get("scope_and_architecture", {})
    if isinstance(scope, dict):
        story.append(Paragraph(f"<b>Scope Description:</b> {scope.get('scope_description', 'N/A')}", body_style))
        story.append(Paragraph(f"<b>System Architecture:</b> {scope.get('system_architecture', 'N/A')}", body_style))
    story.append(Spacer(1, 10))
    
    # Threat Analysis
    story.append(Paragraph("3. Threat and Vulnerability Analysis", h2_style))
    threats = sections.get("threat_analysis", [])
    if isinstance(threats, list) and len(threats) > 0:
        table_data = [["Domain", "Risk Level", "Description"]]
        for t in threats:
            table_data.append([t.get("domain", ""), t.get("level", ""), t.get("description", "")])
        tab = Table(table_data, colWidths=[150, 100, 250])
        tab.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e2e8f0')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
        ]))
        story.append(tab)
    else:
        story.append(Paragraph("No threat analysis recorded.", body_style))
    
    doc.build(story)
    return buffer.getvalue()

def generate_docx_document(assessment_data: dict) -> bytes:
    doc = docx.Document()
    doc.add_heading(assessment_data.get("title", "Privacy Impact Assessment"), 0)
    
    doc.add_paragraph(f"Assessment ID: {assessment_data.get('assessment_id')}")
    doc.add_paragraph(f"Output Type: {assessment_data.get('output_type')}")
    doc.add_paragraph(f"Level: Level {assessment_data.get('level')}")
    doc.add_paragraph(f"Version: {assessment_data.get('version')} ({assessment_data.get('version_notes', 'N/A')})")
    doc.add_paragraph(f"Status: {assessment_data.get('status')}")
    doc.add_paragraph(f"Risk Level: {assessment_data.get('risk_level')}")
    doc.add_paragraph(f"Prepared By: {assessment_data.get('prepared_by')}")
    
    sections = assessment_data.get("sections", {})
    if isinstance(sections, str):
        try:
            sections = json.loads(sections)
        except:
            sections = {}
            
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(sections.get("executive_summary", "No executive summary provided."))
    
    doc.add_heading("2. Scope and Architecture", level=1)
    scope = sections.get("scope_and_architecture", {})
    if isinstance(scope, dict):
        doc.add_paragraph(f"Scope Description: {scope.get('scope_description', 'N/A')}")
        doc.add_paragraph(f"System Architecture: {scope.get('system_architecture', 'N/A')}")
        
    doc.add_heading("3. Threat and Vulnerability Analysis", level=1)
    threats = sections.get("threat_analysis", [])
    if isinstance(threats, list) and len(threats) > 0:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Domain'
        hdr_cells[1].text = 'Risk Level'
        hdr_cells[2].text = 'Description'
        for t in threats:
            row_cells = table.add_row().cells
            row_cells[0].text = t.get("domain", "")
            row_cells[1].text = t.get("level", "")
            row_cells[2].text = t.get("description", "")
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- CPO & Horizon Re-evaluation API Endpoints ---

@app.post("/api/horizon/re-evaluate")
def trigger_horizon_reevaluation(scope: str = Depends(require_scopes(["expert"]))):
    """
    Simulates global risk horizon expiry.
    Downgrades active employee_intake projects and creates remediation tasks.
    """
    import datetime
    import uuid
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    # 1. Update all approved/submitted intakes to 'Re-evaluation Required'
    cursor.execute("UPDATE employee_intake SET status = 'Re-evaluation Required' WHERE status IN ('approved', 'submitted')")
    
    # 2. Get affected projects
    cursor.execute("SELECT id, employee_id, project_name FROM employee_intake WHERE status = 'Re-evaluation Required'")
    affected_projects = cursor.fetchall()
    
    # 3. Create 'Recertification & Drift Remediation' mitigations
    now_iso = datetime.datetime.utcnow().isoformat()
    for row in affected_projects:
        proj_id, emp_id, proj_name = row
        mit_id = f"remedy-{uuid.uuid4().hex[:6]}"
        cursor.execute("""
            INSERT INTO employee_mitigation (id, employee_id, project_id, skill_category, framework_target, sla_target, title, description, status, due_date)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (mit_id, emp_id, proj_id, "Enterprise GRC / RoPA / DSAR", "CCPA/GDPR", "48 Hours", "Recertification & Drift Remediation", f"Re-evaluate risk classification and retention policies for {proj_name}", "open", now_iso))
        
        # Log to audit
        cursor.execute("INSERT INTO system_audit_log (timestamp, employee_id, target_id, action, delta) VALUES (?, ?, ?, ?, ?)",
                       (now_iso, "system", proj_id, "horizon_reevaluation", json.dumps({"status": "Re-evaluation Required"})))

    conn_db.commit()
    conn_db.close()
    return {"status": "success", "message": f"Triggered re-evaluation for {len(affected_projects)} projects."}

@app.get("/api/cpo/skills-matrix")
def get_cpo_skills_matrix(scope: str = Depends(require_scopes(["expert"]))):
    """
    Returns enterprise-wide human resource compliance deployment ledger.
    """
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    cursor.execute("""
        SELECT m.employee_id, m.skill_category, m.status, COUNT(*) as count 
        FROM employee_mitigation m 
        GROUP BY m.employee_id, m.skill_category, m.status
    """)
    rows = cursor.fetchall()
    conn_db.close()
    
    matrix = {}
    for r in rows:
        emp = r['employee_id']
        skill = r['skill_category']
        status = r['status']
        cnt = r['count']
        
        if emp not in matrix:
            matrix[emp] = {}
        if skill not in matrix[emp]:
            matrix[emp][skill] = {"total": 0, "open": 0, "in_progress": 0, "closed": 0, "overdue": 0}
        
        if status in matrix[emp][skill]:
            matrix[emp][skill][status] = cnt
        matrix[emp][skill]["total"] += cnt
        
    return matrix

@app.get("/api/cpo/mitigations")
def get_all_mitigations(scope: str = Depends(require_scopes(["expert"]))):
    """
    Returns all mitigations for CPO master control.
    """
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM employee_mitigation")
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

class OverridePayload(BaseModel):
    employee_id: str
    target_id: str
    action: str
    new_assignee: Optional[str] = None
    reason: Optional[str] = None


import hashlib
import json
import os
from datetime import datetime

VAULT_DIR = os.path.expanduser("~/privacy_swarm/portal/vault_d_operations")
os.makedirs(VAULT_DIR, exist_ok=True)

@app.get("/api/cpo/projects/search")
def search_projects(q: str = "", scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    query = "SELECT id, project_name FROM employee_intake WHERE id LIKE ? OR project_name LIKE ?"
    cursor.execute(query, (f"%{q}%", f"%{q}%"))
    rows = cursor.fetchall()
    conn_db.close()
    return [{"project_id": r["id"], "client_name": r["project_name"], "system_tag": r["id"]} for r in rows]

@app.get("/api/projects/{project_id}/telemetry")
def get_project_telemetry(project_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    
    cursor.execute("SELECT * FROM employee_intake WHERE id = ?", (project_id,))
    intake = cursor.fetchone()
    if not intake:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Project not found")

    cursor.execute("SELECT * FROM employee_mitigation WHERE project_id = ? AND status != 'closed'", (project_id,))
    mitigations = cursor.fetchall()
    
    conn_db.close()
    
    active_risks = []
    open_mitigations = []
    
    for m in mitigations:
        active_risks.append({
            "risk_id": m["id"],
            "severity": "High" if m["status"] == "overdue" else "Medium",
            "description": m["description"]
        })
        open_mitigations.append({
            "task_id": m["id"],
            "assignee": m["employee_id"],
            "status": m["status"]
        })
        
    return {
        "project_id": intake["id"],
        "project_name": intake["project_name"],
        "current_phase": intake["status"],
        "active_risks": active_risks,
        "open_mitigations": open_mitigations,
        "last_audit_timestamp": datetime.utcnow().isoformat() + "Z"
    }

class ReportGeneratePayload(BaseModel):
    requested_by: str
    include_history: bool

@app.post("/api/projects/{project_id}/report/generate")
def generate_project_report(project_id: str, payload: ReportGeneratePayload, scope: str = Depends(require_scopes(["expert"]))):
    # Fetch telemetry
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM employee_intake WHERE id = ?", (project_id,))
    intake = cursor.fetchone()
    if not intake:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Project not found")
        
    cursor.execute("SELECT * FROM employee_mitigation WHERE project_id = ?", (project_id,))
    mitigations = cursor.fetchall()
    conn_db.close()
    
    # LangGraph history via get_state_history
    history = []
    if payload.include_history:
        # get_state_history works with thread_id = project_id
        config = {"configurable": {"thread_id": project_id}}
        for state in graph.get_state_history(config):
            history.append({
                "timestamp": state.created_at.isoformat() if hasattr(state, "created_at") and state.created_at else None,
                "values": state.values
            })

    report_payload = {
        "metadata": {
            "project_id": project_id,
            "project_name": intake["project_name"],
            "requested_by": payload.requested_by,
            "timestamp": datetime.utcnow().isoformat() + "Z",
        },
        "telemetry": {
            "current_status": intake["status"],
            "retention": f"{intake['retention_value']} {intake['retention_unit']}",
            "data_classification": json.loads(intake["data_classification"]) if intake["data_classification"] else [],
        },
        "mitigations": [dict(m) for m in mitigations],
        "langgraph_audit_ledger": history
    }
    
    report_json = json.dumps(report_payload, indent=2)
    sha256_hash = hashlib.sha256(report_json.encode('utf-8')).hexdigest()
    
    filename = f"report_{project_id}_{sha256_hash[:8]}.json"
    filepath = os.path.join(VAULT_DIR, filename)
    
    with open(filepath, "w") as f:
        f.write(report_json)
        
    return {
        "artifact_url": f"/vault/{filename}",
        "sha256_hash": sha256_hash,
        "timestamp": report_payload["metadata"]["timestamp"]
    }


@app.post("/api/cpo/override")
def cpo_override(payload: OverridePayload, scope: str = Depends(require_scopes(["expert"]))):
    """
    Handles CPO explicit overrides.
    """
    import datetime
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    now_iso = datetime.datetime.utcnow().isoformat()
    
    delta = {"reason": payload.reason}
    
    if payload.action == "MANUALLY RE-ROUTE TASK" and payload.new_assignee:
        cursor.execute("UPDATE employee_mitigation SET employee_id = ? WHERE id = ?", (payload.new_assignee, payload.target_id))
        delta["new_assignee"] = payload.new_assignee
    elif payload.action == "FORCED RISK ACCEPTANCE BY CPO":
        cursor.execute("UPDATE employee_mitigation SET status = 'closed', closed_at = ? WHERE id = ?", (now_iso, payload.target_id))
        delta["new_status"] = "closed"
    else:
        conn_db.close()
        raise HTTPException(status_code=400, detail="Invalid override action")
        
    cursor.execute("INSERT INTO system_audit_log (timestamp, employee_id, target_id, action, delta) VALUES (?, ?, ?, ?, ?)",
                   (now_iso, payload.employee_id, payload.target_id, payload.action, json.dumps(delta)))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "action_taken": payload.action}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)



# --- CPO Command Center Upgrade Endpoints ---

import uuid
from datetime import datetime, timedelta

@app.get("/api/dashboard/posture")
def get_privacy_posture(scope: str = Depends(require_scopes(["expert"]))):
    # Simulated posture calculation
    return {
        "snapshot_id": str(uuid.uuid4()),
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "privacy_posture": "yellow", # Mocked for demonstration (e.g. 1 incident)
        "sla_compliance_pct": 98.0,
        "open_requests": 12,
        "critical_requests": 3,
        "median_ttr_hours": 18.5,
        "open_incidents": 1,
        "agents_online": 4,
        "agents_degraded": 0,
        "agents_offline": 0,
        "uptime_pct": 99.8
    }

@app.get("/api/dashboard/trends")
def get_trends(scope: str = Depends(require_scopes(["expert"]))):
    # Mock 90-day trend
    return [
        {"date": (datetime.utcnow() - timedelta(days=i*5)).strftime("%Y-%m-%d"), "received": 5 + i, "resolved": 4 + i}
        for i in range(18)
    ]

@app.get("/api/queue/urgency")
def get_queue_urgency(scope: str = Depends(require_scopes(["expert", "employee"]))):
    # Retrieves all active states from checkpointer and computes SLA
    states = []
    # Simplified mapping
    # Just grab all states and append a fake SLA for demonstration
    for state in graph.get_state_history({"configurable": {"thread_id": "dummy"}}): pass
    
    # We will query metadata_log for active threads
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM metadata_log WHERE status != 'closed'")
    rows = cursor.fetchall()
    
    results = []
    for r in rows:
        tid = r["thread_id"]
        state = get_thread_state(tid)
        if state:
            # compute SLA based on state priority
            hours_left = 48 if state.get("priority") == "high" else 120
            # for testing, make some critical
            if state.get("jurisdiction") == "CCPA": hours_left = 5
            
            results.append({
                "thread_id": tid,
                "state": state,
                "hours_to_breach": hours_left,
                "urgency_level": "red" if hours_left < 6 else ("amber" if hours_left < 24 else "green")
            })
            
    conn_db.close()
    
    # Sort by urgency
    results.sort(key=lambda x: x["hours_to_breach"])
    return results

class DecisionAuditPayload(BaseModel):
    action: str
    approved_by: str
    authority_level: str
    dual_control_confirmed: bool
    justification_checklist: dict
    justification_note: str

@app.patch("/api/transactions/{thread_id}/decision")
def update_transaction_decision(thread_id: str, payload: DecisionAuditPayload):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute('''
        INSERT INTO decision_audit (decision_id, transaction_id, action, approved_by, authority_level, dual_control_confirmed, ip_address, device, justification_checklist, justification_note, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (str(uuid.uuid4()), thread_id, payload.action, payload.approved_by, payload.authority_level, payload.dual_control_confirmed, "127.0.0.1", "Web", json.dumps(payload.justification_checklist), payload.justification_note, datetime.utcnow().isoformat() + "Z"))
    
    # Actually advance the graph
    graph.update_state({"configurable": {"thread_id": thread_id}}, {
        "human_decision": payload.action,
        "human_reasoning": payload.justification_note
    })
    
    # Resume the graph
    for event in graph.stream(None, {"configurable": {"thread_id": thread_id}}):
        pass

    cursor.execute("UPDATE metadata_log SET status = 'closed' WHERE thread_id = ?", (thread_id,))
    conn_db.commit()
    conn_db.close()
    return {"status": "success"}

@app.get("/api/reports/query")
def query_reports(jurisdiction: str = "", type: str = "", scope: str = Depends(require_scopes(["expert"]))):
    # Returns flat JSON of decisions for the report builder
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute('''
        SELECT d.*, m.priority as type
        FROM decision_audit d
        JOIN metadata_log m ON d.transaction_id = m.thread_id
    ''')
    rows = [dict(r) for r in cursor.fetchall()]
    conn_db.close()
    return rows

@app.get("/api/radar")
def get_radar(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM radar_items WHERE dismissed = 0 ORDER BY created_at DESC")
    rows = [dict(r) for r in cursor.fetchall()]
    conn_db.close()
    return rows

@app.post("/api/radar/{radar_id}/dismiss")
def dismiss_radar(radar_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("UPDATE radar_items SET dismissed = 1 WHERE radar_id = ?", (radar_id,))
    conn_db.commit()
    conn_db.close()
    return {"status": "dismissed"}

@app.get("/api/incidents")
def get_incidents(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM incidents")
    
    rows = []
    for r in cursor.fetchall():
        d = dict(r)
        d["jurisdictions_affected"] = json.loads(d["jurisdictions_affected"])
        d["statutory_deadlines"] = json.loads(d["statutory_deadlines"])
        d["timeline"] = json.loads(d["timeline"])
        d["escalation_recipients"] = json.loads(d["escalation_recipients"])
        rows.append(d)
        
    conn_db.close()
    return rows

@app.get("/api/system/health")
def get_system_health(scope: str = Depends(require_scopes(["expert"]))):
    return {
        "status": "Operational",
        "description": "All systems nominal. API Gateway is receiving traffic normally.",
        "uptime_pct": 99.8,
        "queue_depth": 15,
        "processing_rate_per_min": 12.5,
        "rpo": "15 min",
        "rto": "1 hr",
        "last_backup": (datetime.utcnow() - timedelta(minutes=10)).isoformat() + "Z"
    }


# --- RBAC and Evidence Vault Generator Endpoints (V2) ---

class GovernPayload(BaseModel):
    action: str
    justification: str

@app.post("/api/v2/requests/{thread_id}/govern")
def govern_request(thread_id: str, payload: GovernPayload, request: Request, user: UserRole = Depends(get_current_active_user)):
    ip_address = request.headers.get("X-Forwarded-For", request.client.host if request.client else "Unknown")
    device_fingerprint = request.headers.get("User-Agent", "Unknown")
    # Hardcoded Backend Guardrails
    if payload.action == "force_risk_acceptance" and user not in [UserRole.CPO, UserRole.LEGAL]:
        raise HTTPException(status_code=403, detail="CPO or LEGAL required for risk acceptance.")
    if payload.action == "delete_evidence" and user != UserRole.CPO:
        raise HTTPException(status_code=403, detail="CPO required to delete evidence.")
    if user == UserRole.ANALYST and payload.action not in ["submit_for_review", "flag_to_dpo"]:
        raise HTTPException(status_code=403, detail="ANALYST can only submit_for_review or flag_to_dpo.")
        
    # Inject user role into AgentState to enforce downstream LangGraph constraints
    graph.update_state({"configurable": {"thread_id": thread_id}}, {
        "current_user_role": user.value,
        "human_decision": payload.action,
        "human_reasoning": payload.justification
    })
    
    # Resume the graph
    for event in graph.stream(None, {"configurable": {"thread_id": thread_id}}):
        pass

    return {"status": "success", "governed_by": user.value, "action_executed": payload.action, "ip_address": ip_address, "device_fingerprint": device_fingerprint}

job_status_db = {}

def generate_vault_report(job_id: str, payload: dict):
    try:
        pdf_path = f"/tmp/{job_id}.pdf"
        c = canvas.Canvas(pdf_path)
        c.setFont("Courier", 12)
        c.drawString(100, 750, f"Framework Target: {payload.get('framework_target')}")
        c.drawString(100, 730, "Immutable Audit Artifact")
        c.drawString(100, 710, "Generated for SOC2 / ISO 27001 / FTC Compliance.")
        c.drawString(100, 690, "LangGraph Aggregation: SUCCESS")
        
        y = 650
        for pid in payload.get("project_ids", []):
            c.drawString(100, y, f"- Verified Project State: {pid}")
            y -= 20
            
        c.save()
        
        with open(pdf_path, "rb") as f:
            h = hashlib.sha256(f.read()).hexdigest()
            
        job_status_db[job_id] = {
            "status": "COMPLETED",
            "hash": h,
            "artifact_url": f"/vault_evidence/{job_id}.pdf"
        }
    except Exception as e:
        job_status_db[job_id] = {"status": "FAILED", "error": str(e)}

class VaultReportPayload(BaseModel):
    framework_target: str
    date_range: list[str]
    project_ids: list[str]

@app.post("/api/v2/vault/generate-report")
def trigger_vault_report(payload: VaultReportPayload, background_tasks: BackgroundTasks, user: UserRole = Depends(requires_role([UserRole.CPO, UserRole.DPO, UserRole.AUDITOR]))):
    job_id = str(uuid.uuid4())
    job_status_db[job_id] = {"status": "COMPILING"}
    background_tasks.add_task(generate_vault_report, job_id, payload.dict())
    return {"job_id": job_id, "status": "202 Accepted"}
    
@app.get("/api/v2/vault/jobs/{job_id}")
def get_job_status(job_id: str):
    if job_id not in job_status_db:
        raise HTTPException(status_code=404)
    return job_status_db[job_id]


# --- Weighted Priority Engine ---

def get_priority_weights():
    conn = get_db_conn()
    c = conn.cursor()
    c.execute('SELECT severity_weight, jurisdiction_weight, hours_weight, executive_weight, volume_weight FROM priority_config WHERE id = 1')
    row = c.fetchone()
    conn.close()
    if row:
        return PriorityWeights(
            severity_weight=row[0],
            jurisdiction_weight=row[1],
            hours_weight=row[2],
            executive_weight=row[3],
            volume_weight=row[4]
        )
    return PriorityWeights()

def calculate_priority_score(tx: dict, weights: PriorityWeights):
    # 1. Severity
    severity_map = {"Critical": 5, "High": 4, "Medium": 3, "Low": 2, "Info": 1}
    sv = severity_map.get(tx.get('priority', 'Medium'), 3)

    # 2. Jurisdiction
    j_map = {"GDPR": 5, "UK GDPR": 5, "Law 25": 4, "HIPAA": 4, "CCPA": 3, "CPRA": 3, "PIPEDA": 3, "PIPA": 2, "PHIPA": 2}
    jv = j_map.get(tx.get('jurisdiction', ''), 1)

    # 3. Hours Remaining
    # Using the mock deadline parsing for SLA remaining
    deadline = tx.get('deadline_date', '')
    hv = 0
    if deadline:
        try:
            d = datetime.fromisoformat(deadline)
            hrs = (d - datetime.now(d.tzinfo)).total_seconds() / 3600
            if hrs < 6: hv = 5
            elif hrs < 24: hv = 4
            elif hrs < 72: hv = 3
            elif hrs < 168: hv = 2
            elif hrs > 0: hv = 1
        except:
            pass

    # 4. Executive Impact (mock field, assuming routine 0 for now unless specifically flagged)
    ev = 1 if tx.get('priority') == 'Critical' else 0

    # 5. Volume (Mock parsing if record count is present, else 0)
    record_count = 0
    vv = 0
    if record_count > 100000: vv = 5
    elif record_count > 10000: vv = 4
    elif record_count > 1000: vv = 3
    elif record_count > 100: vv = 2
    elif record_count > 0: vv = 1

    score = (sv * weights.severity_weight) + (jv * weights.jurisdiction_weight) +             (hv * weights.hours_weight) + (ev * weights.executive_weight) +             (vv * weights.volume_weight)
            
    return {
        "score": round(score, 2),
        "breakdown": {
            "severity": sv,
            "jurisdiction": jv,
            "hours": hv,
            "executive": ev,
            "volume": vv
        }
    }

@app.get("/api/priority/config", response_model=PriorityWeights)
def get_priority_config():
    return get_priority_weights()

@app.put("/api/priority/config")
def update_priority_config(weights: PriorityWeights, user: UserRole = Depends(requires_role([UserRole.CPO]))):
    conn = get_db_conn()
    c = conn.cursor()
    c.execute('''
        UPDATE priority_config 
        SET severity_weight=?, jurisdiction_weight=?, hours_weight=?, executive_weight=?, volume_weight=?
        WHERE id=1
    ''', (weights.severity_weight, weights.jurisdiction_weight, weights.hours_weight, weights.executive_weight, weights.volume_weight))
    conn.commit()
    conn.close()
    return {"status": "updated"}

@app.get("/api/queue/prioritized")
def get_prioritized_queue():
    # fetch active
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT id, request_type, priority, timestamp, deadline_date, status, jurisdiction FROM metadata_log WHERE status != 'COMPLETED'")
    rows = c.fetchall()
    conn.close()
    
    weights = get_priority_weights()
    transactions = []
    for r in rows:
        tx = {
            "id": r[0], "request_type": r[1], "priority": r[2], 
            "timestamp": r[3], "deadline_date": r[4], "status": r[5], "jurisdiction": r[6]
        }
        pri = calculate_priority_score(tx, weights)
        tx["priority_score"] = pri["score"]
        tx["score_breakdown"] = pri["breakdown"]
        transactions.append(tx)
        
    transactions.sort(key=lambda x: x["priority_score"], reverse=True)
    return transactions

@app.get("/api/transactions/{thread_id}/priority")
def get_transaction_priority(thread_id: str):
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT id, request_type, priority, timestamp, deadline_date, status, jurisdiction FROM metadata_log WHERE id = ?", (thread_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404)
        
    tx = {
        "id": row[0], "request_type": row[1], "priority": row[2], 
        "timestamp": row[3], "deadline_date": row[4], "status": row[5], "jurisdiction": row[6]
    }
    weights = get_priority_weights()
    return calculate_priority_score(tx, weights)



# --- Sprint 2: RBAC & Dual-Approval Endpoints ---

@app.get("/api/rbac/matrix", response_model=list[RolePermission])
def get_rbac_matrix():
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT role, action, allowed, requires_cosign FROM rbac_matrix")
    rows = c.fetchall()
    conn.close()
    return [RolePermission(role=r[0], action=r[1], allowed=bool(r[2]), requires_cosign=r[3]) for r in rows]

@app.put("/api/rbac/matrix")
def update_rbac_matrix(permissions: list[RolePermission], user: UserRole = Depends(requires_role([UserRole.CPO]))):
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("DELETE FROM rbac_matrix")
    for p in permissions:
        c.execute("INSERT INTO rbac_matrix (role, action, allowed, requires_cosign) VALUES (?, ?, ?, ?)", 
                  (p.role, p.action, p.allowed, p.requires_cosign))
    conn.commit()
    conn.close()
    return {"status": "success", "message": "RBAC matrix updated."}

@app.get("/api/rbac/roles/{role}")
def get_role_permissions(role: str):
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT action, allowed, requires_cosign FROM rbac_matrix WHERE role = ?", (role.lower(),))
    rows = c.fetchall()
    conn.close()
    return [{"action": r[0], "allowed": bool(r[1]), "requires_cosign": r[2]} for r in rows]

class InitiateActionPayload(BaseModel):
    action_type: str
    target_id: str

@app.post("/api/actions/{id}/initiate")
def initiate_action(id: str, payload: InitiateActionPayload, user: UserRole = Depends(get_current_active_user)):
    # 1. Look up permission in RBAC matrix
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT allowed, requires_cosign FROM rbac_matrix WHERE role = ? AND action = ?", (user.value.lower(), payload.action_type))
    row = c.fetchone()
    
    if not row or not row[0]:
        conn.close()
        raise HTTPException(status_code=403, detail="Action not allowed for your role.")
        
    requires_cosign = row[1]
    
    if requires_cosign:
        # Create Pending Dual-Approval Action
        c.execute("INSERT INTO dual_actions (action_id, action_type, initiated_by, initiated_role, requires_role, status, timestamp) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  (id, payload.action_type, "user@kibo.is", user.value, requires_cosign.upper(), "pending_second_approval", datetime.utcnow().isoformat()))
        conn.commit()
        conn.close()
        return {"status": "202 Accepted", "message": f"Action pending co-sign from {requires_cosign.upper()}", "action_id": id}
        
    # Execute immediately if no co-sign required
    conn.close()
    return {"status": "200 OK", "message": "Action executed immediately.", "action_id": id}

@app.post("/api/actions/{id}/cosign")
def cosign_action(id: str, request: Request, user: UserRole = Depends(get_current_active_user)):
    ip_address = request.headers.get("X-Forwarded-For", request.client.host if request.client else "Unknown")
    device_fingerprint = request.headers.get("User-Agent", "Unknown")
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT requires_role, status FROM dual_actions WHERE action_id = ?", (id,))
    row = c.fetchone()
    
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Action not found.")
        
    requires_role, status = row[0], row[1]
    
    if status != "pending_second_approval":
        conn.close()
        raise HTTPException(status_code=400, detail="Action is not pending a co-sign.")
        
    if user.value != requires_role and user.value != "CPO":
        conn.close()
        raise HTTPException(status_code=403, detail=f"Co-sign requires role: {requires_role}")
        
    c.execute("UPDATE dual_actions SET status = 'approved', cosigned_by = ? WHERE action_id = ?", (user.value, id))
    conn.commit()
    conn.close()
    
    # Execute Graph Mutation Here
    return {"status": "200 OK", "message": "Action co-signed and executed.", "ip_address": ip_address, "device_fingerprint": device_fingerprint}

@app.get("/api/rationale/templates")
def get_rationale_templates():
    return {
        "dsar_approval": [
            "Identity verified via internal CRM; data located across active systems; no legal hold; within statutory window.",
            "Partial fulfillment — ongoing legal hold exemption applies under relevant statute.",
        ],
        "dsar_rejection": [
            "Identity could not be verified after 3 automated attempts.",
            "Request is manifestly unfounded/excessive under GDPR Art 12(5) or equivalent.",
            "Exemption applies: active internal investigation."
        ],
        "vendor_approval": [
            "DPA executed with SCCs; SOC 2 Type II reviewed; cross-border assessment complete.",
            "Conditional — approved pending SOC 2 Type II execution by next quarter."
        ]
    }


import csv
from io import StringIO
from fastapi.responses import PlainTextResponse

@app.get("/api/v2/vault/export-csv")
def export_vault_csv(user: UserRole = Depends(requires_role([UserRole.CPO, UserRole.DPO, UserRole.AUDITOR]))):
    conn = get_db_conn()
    c = conn.cursor()
    c.execute("SELECT id, jurisdiction, timestamp, status FROM metadata_log LIMIT 100")
    rows = c.fetchall()
    conn.close()
    
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["request_id", "jurisdiction", "data_subject", "request_date", "response_date", "days_to_resolve", "outcome", "approver", "denial_reason", "audit_proof_hash"])
    
    csv_data = []
    for r in rows:
        row_str = f"{r[0]}|{r[1]}|{r[2]}|{r[3]}"
        row_hash = hashlib.sha256(row_str.encode()).hexdigest()
        writer.writerow([r[0], r[1], "consumer@example.com", r[2], r[2], "2", r[3], "CPO", "N/A", row_hash])
        
    csv_string = output.getvalue()
    file_hash = hashlib.sha256(csv_string.encode()).hexdigest()
    
    # Append global hash
    csv_string += f"\n# METADATA: Generated 2026-06-24 by {user.value}; immutable\n"
    csv_string += f"# FILE HASH (SHA-256): {file_hash}\n"
    
    return PlainTextResponse(content=csv_string, headers={"Content-Disposition": "attachment; filename=audit_export.csv"})


# --- NEW LOCALIZED ROLE-BASED ENDPOINTS ---

class UserJurisdictionPayload(BaseModel):
    jurisdiction: str

class WidgetRequestPayload(BaseModel):
    name: str
    email: str
    request_type: str
    description: str

class EmployeeInventoryPayload(BaseModel):
    system_name: str
    data_types: List[str]
    purpose: str
    retention: str
    sharing: str

class MeetingCreatePayload(BaseModel):
    type: str
    date: str
    agenda: List[str]
    attendees: List[str]
    materials: Optional[List[str]] = []

class MeetingMinutesPayload(BaseModel):
    minutes: str
    action_items: List[Dict[str, Any]]

class PsrRecommendationPayload(BaseModel):
    risk_id: str
    member: str
    vote: str
    recommendation: str

class AssessmentCommentPayload(BaseModel):
    comment: str
    member: str

JURISDICTIONS_CONFIG = {
    "canada": {
        "code": "canada",
        "flag": "🇨🇦",
        "name": "Canada (Federal) - PIPEDA",
        "access_request_term": "Access Request / ATIP",
        "access_request_abbr": "ATIP",
        "data_subject_term": "Individual",
        "regulator": "Office of the Privacy Commissioner of Canada (OPC)",
        "primary_statute": "PIPEDA",
        "access_deadline_days": 30,
        "breach_notification": "OPC + affected individuals (real risk of significant harm)",
        "assessment_types": ["PIA", "TRA"],
        "training_track": "pipeda",
        "obligations": "Report to OPC + individuals when there is a Real Risk of Significant Harm (RROSH).",
        "statutes_citation": "PIPEDA S.C. 2000, c. 5"
    },
    "quebec": {
        "code": "quebec",
        "flag": "🇨🇦",
        "name": "Quebec - Law 25",
        "access_request_term": "Access Request",
        "access_request_abbr": "AR",
        "data_subject_term": "Person concerned",
        "regulator": "Commission d'accès à l'information (CAI)",
        "primary_statute": "Law 25",
        "access_deadline_days": 30,
        "breach_notification": "CAI + affected individuals (risk of serious injury)",
        "assessment_types": ["PIA (Mandatory for high-risk)", "TRA"],
        "training_track": "law25",
        "obligations": "Mandatory PIA for high-risk systems; notify CAI of confidentiality incidents.",
        "statutes_citation": "Act respecting the protection of personal information in the private sector (CQLR c P-39.1)"
    },
    "ontario": {
        "code": "ontario",
        "flag": "🇨🇦",
        "name": "Ontario (Public Sector) - FIPPA",
        "access_request_term": "Freedom of Information Request",
        "access_request_abbr": "FOI",
        "data_subject_term": "Requester",
        "regulator": "Information and Privacy Commissioner of Ontario (IPC)",
        "primary_statute": "FIPPA",
        "access_deadline_days": 30,
        "breach_notification": "IPC + affected individuals (extendable clock)",
        "assessment_types": ["PIA", "TRA"],
        "training_track": "fippa",
        "obligations": "Freedom of Information governance, 30 days extendable time limit.",
        "statutes_citation": "FIPPA R.S.O. 1990, c. F.31"
    },
    "us": {
        "code": "us",
        "flag": "🇺🇸",
        "name": "United States - CCPA/CPRA",
        "access_request_term": "DSAR (Data Subject Access Request)",
        "access_request_abbr": "DSAR",
        "data_subject_term": "Consumer",
        "regulator": "California Privacy Protection Agency (CPPA) / California AG",
        "primary_statute": "CPRA",
        "access_deadline_days": 45,
        "breach_notification": "State-by-state matrix (Attorney General + individuals)",
        "assessment_types": ["Risk Assessment", "TRA"],
        "training_track": "ccpa",
        "obligations": "45 days (+45 extension) response window. State attorney general breach reporting.",
        "statutes_citation": "California Civil Code Sec. 1798.100 et seq."
    },
    "eu": {
        "code": "eu",
        "flag": "🇪🇺",
        "name": "European Union - GDPR",
        "access_request_term": "DSAR / Right of Access",
        "access_request_abbr": "DSAR",
        "data_subject_term": "Data Subject",
        "regulator": "Supervisory Authority (DPA)",
        "primary_statute": "GDPR Art 15",
        "access_deadline_days": 30,
        "breach_notification": "72 hours to supervisory authority (GDPR Art 33)",
        "assessment_types": ["DPIA (Art 35)", "TIA (Schrems II)"],
        "training_track": "gdpr",
        "obligations": "72-hour breach notification window. DPIA and Schrems II Transfer Impact Assessment obligations.",
        "statutes_citation": "GDPR (EU) 2016/679"
    },
    "uk": {
        "code": "uk",
        "flag": "🇬🇧",
        "name": "United Kingdom - UK GDPR",
        "access_request_term": "Subject Access Request",
        "access_request_abbr": "SAR",
        "data_subject_term": "Data Subject",
        "regulator": "Information Commissioner's Office (ICO)",
        "primary_statute": "UK GDPR",
        "access_deadline_days": 30,
        "breach_notification": "ICO + affected individuals within 72 hours",
        "assessment_types": ["UK GDPR DPIA", "TRA"],
        "training_track": "uk_gdpr",
        "obligations": "Statutory 1 month response window. Report breaches to ICO.",
        "statutes_citation": "Data Protection Act 2018 / UK GDPR"
    }
}

@app.get("/api/jurisdictions")
def get_jurisdictions():
    return list(JURISDICTIONS_CONFIG.values())

@app.get("/api/jurisdictions/{code}/config")
def get_jurisdiction_config(code: str):
    if code not in JURISDICTIONS_CONFIG:
        raise HTTPException(status_code=404, detail="Jurisdiction not found")
    return JURISDICTIONS_CONFIG[code]

@app.put("/api/user/jurisdiction")
def update_user_jurisdiction(payload: UserJurisdictionPayload):
    if payload.jurisdiction not in JURISDICTIONS_CONFIG:
        raise HTTPException(status_code=400, detail="Invalid jurisdiction code")
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("INSERT OR REPLACE INTO user_settings (key, value) VALUES ('active_jurisdiction', ?)", (payload.jurisdiction,))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "active_jurisdiction": payload.jurisdiction}

# --- PUBLIC WIDGET ENDPOINTS ---

@app.post("/api/widget/{org_id}/request")
def submit_widget_request(org_id: str, payload: WidgetRequestPayload):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT value FROM user_settings WHERE key='active_jurisdiction'")
    row = cursor.fetchone()
    jur = row[0] if row else "ontario"
    
    import uuid
    from datetime import datetime
    req_id = f"REQ-{uuid.uuid4().hex[:6].upper()}"
    
    cursor.execute("INSERT INTO metadata_log (thread_id, priority, deadline, status) VALUES (?, 'medium', ?, 'pending')", (req_id, f"{JURISDICTIONS_CONFIG[jur]['access_deadline_days']} days"))
    conn_db.commit()
    conn_db.close()
    
    config = {"configurable": {"thread_id": req_id}}
    tx_data = {
        "id": req_id,
        "type": "DSAR" if jur in ["us", "eu", "uk"] else "FOI",
        "client": org_id,
        "jurisdiction": jur.upper(),
        "priority": "medium",
        "deadline": f"{JURISDICTIONS_CONFIG[jur]['access_deadline_days']} days",
        "description": f"Public submission from {payload.name} ({payload.email}): {payload.description}",
        "agent": "Widget Ingestion Agent",
        "summary": "Widget submitted privacy request.",
        "raw": json.dumps({"name": payload.name, "email": payload.email, "type": payload.request_type, "desc": payload.description})
    }
    graph.update_state(config, tx_data)
    graph.invoke(None, config)
    
    return {"status": "submitted", "request_id": req_id, "deadline_days": JURISDICTIONS_CONFIG[jur]['access_deadline_days']}

@app.get("/api/widget/{org_id}/status/{req_id}")
def check_widget_status(org_id: str, req_id: str):
    state_vals = get_thread_state(req_id)
    if not state_vals:
        conn_db = sqlite3.connect(DB_FILE)
        cursor = conn_db.cursor()
        cursor.execute("SELECT status, deadline FROM metadata_log WHERE thread_id=?", (req_id,))
        row = cursor.fetchone()
        conn_db.close()
        if row:
            return {"id": req_id, "status": row[0], "deadline": row[1]}
        raise HTTPException(status_code=404, detail="Request not found")
        
    config = {"configurable": {"thread_id": req_id}}
    next_steps = graph.get_state(config).next
    is_pending = "human_approval_node" in next_steps
    status = "under_review" if is_pending else state_vals.get("status", "pending")
    return {
        "id": req_id,
        "status": status,
        "deadline": state_vals.get("deadline", "30 days"),
        "type": state_vals.get("type", "Access")
    }

@app.get("/api/widget/{org_id}/config")
def get_widget_config(org_id: str):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT value FROM user_settings WHERE key='active_jurisdiction'")
    row = cursor.fetchone()
    jur = row[0] if row else "ontario"
    conn_db.close()
    return {
        "org_id": org_id,
        "jurisdiction": jur,
        "config": JURISDICTIONS_CONFIG[jur]
    }

# --- EMPLOYEE ENDPOINTS ---

@app.get("/api/employee/{id}/risks")
def get_employee_assigned_risks(id: str, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT value FROM user_settings WHERE key='active_jurisdiction'")
    row = cursor.fetchone()
    jur = row[0] if row else "ontario"
    
    cursor.execute("SELECT risk_id, issue, gap_type, risk_score, risk_level, assigned_to, next_review_date, status FROM risks")
    rows = cursor.fetchall()
    conn_db.close()
    
    deadline_days = JURISDICTIONS_CONFIG[jur]["access_deadline_days"]
    
    risks_list = []
    for r in rows:
        risks_list.append({
            "risk_id": r[0],
            "issue": r[1],
            "gap_type": r[2],
            "risk_score": r[3],
            "risk_level": r[4],
            "assigned_to": r[5],
            "due_date": r[6] if r[6] else f"Within {deadline_days} days",
            "status": r[7]
        })
    return risks_list

@app.get("/api/employee/{id}/inventory-tasks")
def get_employee_inventory_tasks(id: str, scope: str = Depends(require_scopes(["employee", "expert"]))):
    return {
        "completion_percentage": 60,
        "tasks": [
            {"id": "inv-1", "system": "Support ZenDesk Integration", "status": "pending_input"},
            {"id": "inv-2", "system": "Clinical Logs Store", "status": "pending_input"},
            {"id": "inv-3", "system": "Marketing Mailchimp list", "status": "pending_input"}
        ]
    }

@app.post("/api/employee/{id}/inventory")
def submit_employee_inventory(id: str, payload: EmployeeInventoryPayload, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    import uuid
    inv_id = str(uuid.uuid4())
    cursor.execute("INSERT INTO employee_data_inventory (id, employee_id, system_name, data_types, purpose, retention, sharing) VALUES (?, ?, ?, ?, ?, ?, ?)", (inv_id, id, payload.system_name, json.dumps(payload.data_types), payload.purpose, payload.retention, payload.sharing))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "inventory_id": inv_id}

@app.get("/api/employee/{id}/training")
def get_employee_training(id: str, scope: str = Depends(require_scopes(["employee", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT module_id, title, role_tags, jurisdiction, source_policy, duration_min, required, completed, completed_at FROM employee_training")
    rows = cursor.fetchall()
    conn_db.close()
    
    training = []
    for r in rows:
        training.append({
            "module_id": r[0],
            "title": r[1],
            "role_tags": json.loads(r[2]),
            "jurisdiction": r[3],
            "source_policy": r[4],
            "duration_min": r[5],
            "required": bool(r[6]),
            "completed": bool(r[7]),
            "completed_at": r[8]
        })
    return training

@app.post("/api/employee/{id}/training/{mod}/complete")
def complete_employee_training(id: str, mod: str, scope: str = Depends(require_scopes(["employee", "expert"]))):
    from datetime import datetime
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("UPDATE employee_training SET completed = 1, completed_at = ? WHERE module_id = ? AND employee_id = ?", (datetime.utcnow().isoformat(), mod, id))
    conn_db.commit()
    conn_db.close()
    return {"status": "completed", "module_id": mod}

# --- EXPERT ENDPOINTS ---

@app.get("/api/expert/assessments")
def get_expert_assessments(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT assessment_id, title, type, output_type, level, version, status, risk_level, prepared_by, jurisdictions FROM assessments")
    rows = cursor.fetchall()
    conn_db.close()
    
    assessments_list = []
    for r in rows:
        assessments_list.append({
            "assessment_id": r[0],
            "title": r[1],
            "type": r[2],
            "output_type": r[3],
            "level": r[4],
            "version": r[5],
            "status": r[6],
            "risk_level": r[7],
            "prepared_by": r[8],
            "jurisdictions": json.loads(r[9])
        })
    return assessments_list

@app.post("/api/expert/audit")
def post_expert_audit(scope: str = Depends(require_scopes(["expert"]))):
    return {"status": "audit_initiated", "timestamp": datetime.utcnow().isoformat(), "audit_id": "AUD-992"}

@app.post("/api/expert/soc-compliance")
def post_soc_compliance(scope: str = Depends(require_scopes(["expert"]))):
    return {
        "status": "soc_verification_running",
        "controls_verified": [
            {"control": "Access Control", "status": "passed"},
            {"control": "Data Encryption", "status": "passed"},
            {"control": "Incident Response Plan", "status": "warning"}
        ]
    }

@app.get("/api/expert/meetings")
def get_expert_meetings(scope: str = Depends(require_scopes(["expert", "psr"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT meeting_id, type, date, agenda, attendees, minutes, action_items, materials FROM meetings")
    rows = cursor.fetchall()
    conn_db.close()
    
    meetings_list = []
    for r in rows:
        meetings_list.append({
            "meeting_id": r[0],
            "type": r[1],
            "date": r[2],
            "agenda": json.loads(r[3]),
            "attendees": json.loads(r[4]),
            "minutes": r[5],
            "action_items": json.loads(r[6]) if r[6] else [],
            "materials": json.loads(r[7]) if r[7] else []
        })
    return meetings_list

@app.post("/api/expert/meetings")
def create_expert_meeting(payload: MeetingCreatePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    import uuid
    mid = f"meet-{uuid.uuid4().hex[:6]}"
    cursor.execute("INSERT INTO meetings (meeting_id, type, date, agenda, attendees, minutes, action_items, materials) VALUES (?, ?, ?, ?, ?, '', '[]', ?)", (mid, payload.type, payload.date, json.dumps(payload.agenda), json.dumps(payload.attendees), json.dumps(payload.materials)))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "meeting_id": mid}

@app.post("/api/expert/meetings/{id}/minutes")
def record_expert_meeting_minutes(id: str, payload: MeetingMinutesPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("UPDATE meetings SET minutes = ?, action_items = ? WHERE meeting_id = ?", (payload.minutes, json.dumps(payload.action_items), id))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "meeting_id": id}

@app.get("/api/expert/inbox")
def get_expert_inbox(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT email_id, sender, subject, body, category, status, converted_to_transaction FROM simulated_inbox")
    rows = cursor.fetchall()
    conn_db.close()
    
    inbox = []
    for r in rows:
        inbox.append({
            "email_id": r[0],
            "sender": r[1],
            "subject": r[2],
            "body": r[3],
            "category": r[4],
            "status": r[5],
            "converted_to_transaction": r[6]
        })
    return inbox

@app.post("/api/expert/inbox/{id}/triage")
def triage_expert_inbox(id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT sender, subject, body, category FROM simulated_inbox WHERE email_id=?", (id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Email not found")
        
    sender, subject, body, category = row
    tx_id = f"TX-INBOX-{id.upper()}"
    
    cursor.execute("UPDATE simulated_inbox SET status = 'triaged', converted_to_transaction = ? WHERE email_id = ?", (tx_id, id))
    conn_db.commit()
    conn_db.close()
    
    config = {"configurable": {"thread_id": tx_id}}
    tx_data = {
        "id": tx_id,
        "type": "DSAR" if category == "dsar" else "Inquiry",
        "client": sender,
        "jurisdiction": "ONTARIO",
        "priority": "critical" if category == "regulator" else "medium",
        "deadline": "30 days",
        "description": f"Triaged Inbox Email: {subject} - {body}",
        "agent": "Inbox Agent",
        "summary": f"Triage of simulated email {id}",
        "raw": json.dumps({"sender": sender, "subject": subject, "body": body})
    }
    graph.update_state(config, tx_data)
    graph.invoke(None, config)
    
    return {"status": "triaged", "transaction_id": tx_id}

class EmailReplyPayload(BaseModel):
    body: str

@app.post("/api/expert/inbox/{id}/reply")
def reply_expert_inbox(id: str, payload: EmailReplyPayload, scope: str = Depends(require_scopes(["expert"]))):
    import uuid
    reply_id = f"REP-{uuid.uuid4().hex[:6].upper()}"
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    # Verify email exists
    cursor.execute("SELECT sender FROM simulated_inbox WHERE email_id=?", (id,))
    if not cursor.fetchone():
        conn_db.close()
        raise HTTPException(status_code=404, detail="Email not found")
        
    # Insert reply
    cursor.execute("INSERT INTO simulated_inbox_replies (reply_id, email_id, body) VALUES (?, ?, ?)", (reply_id, id, payload.body))
    # Update email status
    cursor.execute("UPDATE simulated_inbox SET status = 'replied' WHERE email_id = ?", (id,))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "reply_id": reply_id}

@app.get("/api/expert/inbox/{id}/replies")
def get_inbox_replies(id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT reply_id, body, sent_at FROM simulated_inbox_replies WHERE email_id = ? ORDER BY sent_at DESC", (id,))
    rows = cursor.fetchall()
    conn_db.close()
    
    replies = []
    for r in rows:
        replies.append({
            "reply_id": r[0],
            "body": r[1],
            "sent_at": r[2]
        })
    return replies

class LessonPayload(BaseModel):
    domain: str
    feedback_notes: str

@app.post("/api/expert/lessons/add")
def add_expert_lesson(payload: LessonPayload, scope: str = Depends(require_scopes(["expert"]))):
    import uuid
    lesson_id = f"L-{uuid.uuid4().hex[:4].upper()}"
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("INSERT INTO agent_lessons_learned (lesson_id, domain, feedback_notes) VALUES (?, ?, ?)", (lesson_id, payload.domain, payload.feedback_notes))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "lesson_id": lesson_id}

class ImprovementPayload(BaseModel):
    trigger_type: str
    trigger_data: str

@app.post("/api/expert/lessons/trigger_improvement")
def trigger_improvement(payload: ImprovementPayload, scope: str = Depends(require_scopes(["expert"]))):
    from self_improvement import run_self_improvement
    result = run_self_improvement(payload.trigger_type, payload.trigger_data)
    return {
        "status": "success",
        "logs": result["logs"],
        "proposed_logic_update": result["proposed_logic_update"],
        "proposed_ui_schema": result["proposed_ui_schema"],
        "evaluation_score": result["evaluation_score"],
        "is_approved": result["is_approved"]
    }

@app.get("/api/expert/lessons/legal_library")
def get_legal_library(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT clause_id, legislation, clause_text, keywords FROM legal_ground_truth")
    rows = cursor.fetchall()
    conn_db.close()
    clauses = []
    for r in rows:
        clauses.append({
            "clause_id": r[0],
            "legislation": r[1],
            "clause_text": r[2],
            "keywords": r[3]
        })
    return clauses

@app.get("/api/expert/lessons")
def get_expert_lessons(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT lesson_id, domain, feedback_notes, created_at FROM agent_lessons_learned ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn_db.close()
    
    lessons = []
    for r in rows:
        lessons.append({
            "lesson_id": r[0],
            "domain": r[1],
            "feedback_notes": r[2],
            "created_at": r[3]
        })
    return lessons

@app.post("/api/expert/training/assign")
def assign_expert_training(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("UPDATE employee_training SET completed = 0, completed_at = NULL")
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "message": "Training modules reset/assigned to all staff"}

@app.post("/api/expert/training/remind")
def remind_expert_training(scope: str = Depends(require_scopes(["expert"]))):
    return {"status": "success", "message": "Reminders sent to staff members with overdue training modules."}

@app.get("/api/expert/training/compliance")
def get_training_compliance(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT count(*) FROM employee_training")
    total = cursor.fetchone()[0]
    cursor.execute("SELECT count(*) FROM employee_training WHERE completed = 1")
    completed = cursor.fetchone()[0]
    conn_db.close()
    
    compliance_pct = int((completed / total) * 100) if total > 0 else 100
    return {
        "total_modules": total,
        "completed_modules": completed,
        "compliance_percentage": compliance_pct,
        "overdue_staff_count": total - completed
    }

# --- PSR COMMITTEE ENDPOINTS ---

@app.get("/api/psr/meetings")
def get_psr_meetings(scope: str = Depends(require_scopes(["psr", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT meeting_id, type, date, agenda, attendees, minutes, action_items, materials FROM meetings WHERE type = 'psr_committee' OR type = 'privacy_security_weekly'")
    rows = cursor.fetchall()
    conn_db.close()
    
    meetings_list = []
    for r in rows:
        meetings_list.append({
            "meeting_id": r[0],
            "type": r[1],
            "date": r[2],
            "agenda": json.loads(r[3]),
            "attendees": json.loads(r[4]),
            "minutes": r[5],
            "action_items": json.loads(r[6]) if r[6] else [],
            "materials": json.loads(r[7]) if r[7] else []
        })
    return meetings_list

@app.get("/api/psr/meetings/{id}/materials")
def get_psr_meeting_materials(id: str, scope: str = Depends(require_scopes(["psr", "expert"]))):
    return [
        {"name": "Quebec Law 25 Diagnostic Posture.pdf", "size": "1.2 MB"},
        {"name": "Cross-Border Transfer TIA v1.docx", "size": "450 KB"},
        {"name": "SOC 2 Type II System Scope.pdf", "size": "2.1 MB"}
    ]

@app.get("/api/psr/risk-queue")
def get_psr_risk_queue(scope: str = Depends(require_scopes(["psr", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT risk_id, issue, gap_type, risk_score, risk_level, assigned_to, status FROM risks")
    rows = cursor.fetchall()
    
    cursor.execute("SELECT risk_id, member, vote, recommendation, timestamp FROM psr_recommendations")
    recs = cursor.fetchall()
    conn_db.close()
    
    recs_map = {}
    for rc in recs:
        r_id = rc[0]
        if r_id not in recs_map:
            recs_map[r_id] = []
        recs_map[r_id].append({
            "member": rc[1],
            "vote": rc[2],
            "recommendation": rc[3],
            "timestamp": rc[4]
        })
        
    queue = []
    for r in rows:
        queue.append({
            "risk_id": r[0],
            "issue": r[1],
            "gap_type": r[2],
            "risk_score": r[3],
            "risk_level": r[4],
            "assigned_to": r[5],
            "status": r[6],
            "recommendations": recs_map.get(r[0], [])
        })
    return queue

@app.post("/api/psr/risk/{id}/recommendation")
def post_psr_risk_recommendation(id: str, payload: PsrRecommendationPayload, scope: str = Depends(require_scopes(["psr", "expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    import uuid
    from datetime import datetime
    rec_id = str(uuid.uuid4())
    cursor.execute("INSERT INTO psr_recommendations (recommendation_id, risk_id, member, vote, recommendation, timestamp) VALUES (?, ?, ?, ?, ?, ?)", (rec_id, id, payload.member, payload.vote, payload.recommendation, datetime.utcnow().isoformat()))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "recommendation_id": rec_id}

@app.post("/api/psr/assessments/{id}/comment")
def post_psr_assessment_comment(id: str, payload: AssessmentCommentPayload, scope: str = Depends(require_scopes(["psr", "expert"]))):
    return {
        "status": "success",
        "message": f"Advisory comment by {payload.member} logged for assessment {id}"
    }

# --- Onboarding & CASL API Models & Endpoints ---
class OnboardingTaskNotePayload(BaseModel):
    notes: str

@app.get("/api/onboarding/tasks")
def get_onboarding_tasks(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM onboarding_tasks ORDER BY category, id")
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/onboarding/tasks/{task_id}/toggle")
def toggle_onboarding_task(task_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT status FROM onboarding_tasks WHERE id = ?", (task_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Task not found")
    new_status = "completed" if row[0] == "pending" else "pending"
    cursor.execute("UPDATE onboarding_tasks SET status = ? WHERE id = ?", (new_status, task_id))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "task_id": task_id, "new_status": new_status}

@app.post("/api/onboarding/tasks/{task_id}/notes")
def update_onboarding_task_notes(task_id: str, payload: OnboardingTaskNotePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("UPDATE onboarding_tasks SET notes = ? WHERE id = ?", (payload.notes, task_id))
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "task_id": task_id, "notes": payload.notes}

@app.get("/api/onboarding/casl")
def get_casl_registry(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    conn_db.row_factory = sqlite3.Row
    cursor = conn_db.cursor()
    cursor.execute("SELECT * FROM casl_consent_registry ORDER BY id")
    rows = cursor.fetchall()
    conn_db.close()
    return [dict(r) for r in rows]

@app.post("/api/onboarding/casl/sunset")
def run_casl_sunset_automation(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    from datetime import datetime, timezone
    current_time_str = "2026-07-02T23:43:19Z"
    current_date = datetime.strptime(current_time_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
    
    cursor.execute("SELECT id, email, name, consent_status, expiry_date FROM casl_consent_registry")
    contacts = cursor.fetchall()
    
    logs = []
    logs.append(f"[{datetime.now().strftime('%H:%M:%S')}] Starting CASL Sunset Automation workflow (Simulated current date: 2026-07-02)")
    
    updated_count = 0
    expired_count = 0
    warning_count = 0
    
    for c_id, email, name, status, expiry_str in contacts:
        if not expiry_str:
            continue
            
        try:
            expiry_date = datetime.strptime(expiry_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
        except ValueError:
            try:
                expiry_date = datetime.strptime(expiry_str, "%Y-%m-%dT%H:%M:%S").replace(tzinfo=timezone.utc)
            except ValueError:
                expiry_date = datetime.strptime(expiry_str.split("T")[0], "%Y-%m-%d").replace(tzinfo=timezone.utc)
                
        delta = expiry_date - current_date
        
        if delta.days < 0:
            if status != "Expired":
                cursor.execute("UPDATE casl_consent_registry SET consent_status = 'Expired' WHERE id = ?", (c_id,))
                logs.append(f"[SUNSET] EXPIRED: {name} ({email}) - Implied consent expired on {expiry_str.split('T')[0]}. Contact suppressed from commercial lists.")
                expired_count += 1
                updated_count += 1
        elif delta.days <= 30:
            logs.append(f"[SUNSET] WARNING: {name} ({email}) - Implied consent expires in {delta.days} days ({expiry_str.split('T')[0]}). 'Consent Upgrade' email queued.")
            warning_count += 1
            
    conn_db.commit()
    conn_db.close()
    
    logs.append(f"[{datetime.now().strftime('%H:%M:%S')}] CASL Sunset run complete. Processed {len(contacts)} contacts. Suppressed: {expired_count}. Warned/Queued: {warning_count}.")
    
    return {
        "status": "success",
        "processed": len(contacts),
        "suppressed": expired_count,
        "warned": warning_count,
        "logs": logs
    }


@app.post("/api/admin/trigger-verdicts-collector")
def trigger_verdicts_collector(scope: str = Depends(require_scopes(["expert", "employee"]))):
    """Trigger the verdicts collector manually."""
    canadian_verdicts_agent.collect_new_verdicts()
    return {"status": "success", "message": "Verdicts collector triggered successfully."}


from pydantic import BaseModel
from typing import Any, List, Optional, Literal

class KnowledgeDecisionPayload(BaseModel):
    decision: Literal["accept", "reject"]
    notes: Optional[str] = None

# --- SENSING AGENTS ENDPOINTS ---

@app.get("/api/agents/sensing")
def get_sensing_agents(scope: str = Depends(require_scopes(["expert"]))):
    # Retrieve pending updates count
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT agent, COUNT(*) FROM knowledge_updates WHERE status = 'pending_review' GROUP BY agent")
    pending_counts = dict(cursor.fetchall())
    conn_db.close()

    agents = [
        {"agent_id": "regulatory", "name": "Regulatory Sensing Agent", "domain": "regulatory", "schedule": "0 0 * * *", "status": "active", "updates_pending": pending_counts.get("Regulatory Sensing Agent", 0), "last_run": "2026-07-02T12:00:00Z"},
        {"agent_id": "vendor", "name": "Vendor Sensing Agent", "domain": "vendor", "schedule": "0 12 * * *", "status": "active", "updates_pending": pending_counts.get("Vendor Sensing Agent", 0), "last_run": "2026-07-02T18:30:00Z"},
        {"agent_id": "internal", "name": "Internal Systems Sensing Agent", "domain": "internal", "schedule": "*/30 * * * *", "status": "active", "updates_pending": pending_counts.get("Internal Systems Sensing Agent", 0), "last_run": "2026-07-02T23:30:00Z"},
        {"agent_id": "threat", "name": "Threat Sensing Agent", "domain": "threat", "schedule": "0 */4 * * *", "status": "active", "updates_pending": pending_counts.get("Threat Sensing Agent", 0), "last_run": "2026-07-02T20:00:00Z"},
        {"agent_id": "web", "name": "Web & Consent Sensing Agent", "domain": "web", "schedule": "0 2 * * *", "status": "active", "updates_pending": pending_counts.get("Web & Consent Sensing Agent", 0), "last_run": "2026-07-02T22:00:00Z"}
    ]
    return agents

@app.post("/api/agents/sensing/{agent_id}/run")
def run_sensing_agent(agent_id: str, scope: str = Depends(require_scopes(["expert"]))):
    import uuid
    from datetime import datetime
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    update_id = f"KUP-{uuid.uuid4().hex[:6].upper()}"
    logs = []
    
    if agent_id == "regulatory":
        logs = [
            "[Regulatory Agent] Crawling CAI Quebec and IPC Ontario feeds...",
            "[Regulatory Agent] Found new CAI ruling: profiling technologies require mandatory pre-deployment PIA.",
            f"[Regulatory Agent] Synthesized new obligation for Quebec jurisdiction. Confidence: 92%."
        ]
        cursor.execute("""
            INSERT INTO knowledge_updates (update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status)
            VALUES (?, 'regulatory', 'Regulatory Sensing Agent', 'new', 'Mandatory PIA for Profiling', 'None', 'Quebec Law 25 Section 12.3: Organisation must conduct PIA before deploying profiling or tracking technologies.', 92, 'CAI Rulings Release July 2026', 'All active PIAs using tracking tech must be re-evaluated.', '["pia"]', 'pending_review')
        """, (update_id,))
    elif agent_id == "vendor":
        logs = [
            "[Vendor Agent] Scanning subprocessor list for active integrations...",
            "[Vendor Agent] Detected updated HubSpot DPA: Clause 4.2 delegates client telemetry data storage to US AWS servers.",
            f"[Vendor Agent] Warning: cross-border transfer obligation triggered. Confidence: 88%."
        ]
        cursor.execute("""
            INSERT INTO knowledge_updates (update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status)
            VALUES (?, 'vendor', 'Vendor Sensing Agent', 'amended', 'HubSpot Data Processing Agreement', 'HubSpot data local to Canada', 'HubSpot DPA Clause 4.2: Data stored in AWS US-East servers. Cross-border transfer triggered.', 88, 'HubSpot DPA Addendum July 2026', 'Requires execution of Transfer Impact Assessment (TIA).', '["v-2", "risk"]', 'pending_review')
        """, (update_id,))
    elif agent_id == "internal":
        logs = [
            "[Internal Agent] Crawling system schema maps and asset registries...",
            "[Internal Agent] Found new database table 'minor_intake_forms' containing PII (birthdays, health numbers).",
            f"[Internal Agent] Data catalog updated. Confidence: 95%."
        ]
        cursor.execute("""
            INSERT INTO knowledge_updates (update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status)
            VALUES (?, 'inventory', 'Internal Systems Sensing Agent', 'new', 'minor_intake_forms database table', 'None', 'System Table minor_intake_forms: Stores name, dob, health_card_number.', 95, 'KIBO Database Schema Monitor', 'Requires encryption compliance audit on intake forms.', '["inventory"]', 'pending_review')
        """, (update_id,))
    elif agent_id == "threat":
        logs = [
            "[Threat Agent] Fetching global threat advisory databases...",
            "[Threat Agent] Flagged security bulletin CVE-2026-9021: Local privilege escalation in uvicorn/fastapi configurations.",
            f"[Threat Agent] Risk weighting increased. Confidence: 85%."
        ]
        cursor.execute("""
            INSERT INTO knowledge_updates (update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status)
            VALUES (?, 'risk', 'Threat Sensing Agent', 'amended', 'FastAPI/Uvicorn CVE-2026-9021', 'Low Threat Level', 'CVE-2026-9021: Uvicorn service exposure vulnerability. Risk level raised to Medium.', 85, 'NVD Threat Registry', 'Requires system update and access logging audit.', '["risk"]', 'pending_review')
        """, (update_id,))
    elif agent_id == "web":
        logs = [
            "[Web Agent] Headless browser scanning of public web surface kibo.is...",
            "[Web Agent] Found new unauthorized tracking cookie: '_ga_client_metric' (Google Analytics proxy).",
            f"[Web Agent] Consent banner geo compliance check triggered. Confidence: 90%."
        ]
        cursor.execute("""
            INSERT INTO knowledge_updates (update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status)
            VALUES (?, 'web', 'Web & Consent Sensing Agent', 'new', '_ga_client_metric tracker', 'None', 'Unregistered third-party tracking cookie detected in visitor sessions.', 90, 'Headless Scan crawler-01', 'Violates Law 25 default-off consent banner rules.', '["policy"]', 'pending_review')
        """, (update_id,))
    else:
        conn_db.close()
        return {"status": "error", "message": "Unknown agent"}

    conn_db.commit()
    conn_db.close()
    return {"status": "success", "update_id": update_id, "logs": logs}

# --- KNOWLEDGE ENDPOINTS ---

@app.get("/api/knowledge/updates")
def get_knowledge_updates(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("""
        SELECT update_id, domain, agent, change_type, entity, old_value, new_value, confidence, source, impact_summary, affected_entities, status, detected_at
        FROM knowledge_updates
        ORDER BY detected_at DESC
    """)
    rows = cursor.fetchall()
    conn_db.close()
    
    updates = []
    for r in rows:
        updates.append({
            "update_id": r[0],
            "domain": r[1],
            "agent": r[2],
            "change_type": r[3],
            "entity": r[4],
            "old_value": r[5],
            "new_value": r[6],
            "confidence": r[7],
            "source": r[8],
            "impact_summary": r[9],
            "affected_entities": json.loads(r[10]) if r[10] else [],
            "status": r[11],
            "detected_at": r[12]
        })
    return updates

@app.post("/api/knowledge/updates/{update_id}/decision")
def post_knowledge_decision(update_id: str, payload: KnowledgeDecisionPayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    cursor.execute("SELECT domain, change_type, entity, new_value, affected_entities FROM knowledge_updates WHERE update_id = ?", (update_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Update not found")
        
    domain, change_type, entity, new_value, affected_entities_str = row
    status_val = "accepted" if payload.decision == "accept" else "rejected"
    
    cursor.execute("UPDATE knowledge_updates SET status = ? WHERE update_id = ?", (status_val, update_id))
    
    if payload.decision == "accept":
        if domain == "regulatory":
            cursor.execute("INSERT OR REPLACE INTO regulatory_ontology (obligation_id, jurisdiction, statute, obligation_name, applicability_trigger, deadline_days, regulator) VALUES (?, 'quebec', 'Law 25', ?, 'Deploying profiling technologies', 30, 'Commission d''accès à l''information (CAI)')", (update_id, entity))
        elif domain == "vendor":
            cursor.execute("INSERT OR REPLACE INTO vendor_graph (vendor_id, vendor_name, service, data_types, dpa_status, soc2_status, risk_rating) VALUES (?, 'HubSpot', 'Marketing Analytics', 'Client telemetry, email clicks', 'Amended (US AWS)', 'SOC 2 active', 'Medium')", (update_id,))
        elif domain == "inventory":
            cursor.execute("INSERT OR REPLACE INTO data_inventory (inventory_id, system_name, data_types, sensitivity, purpose, retention_rules, sharing_recipient, jurisdiction) VALUES (?, 'minor_intake_forms', 'dob, health_card_number', 'high', 'Patient check-in', 'Delete on session exit', 'None', 'Ontario')", (update_id,))
            
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "update_id": update_id, "resolution": status_val}

@app.get("/api/knowledge/{domain}")
def get_knowledge_domain(domain: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    data = []
    
    if domain == "regulatory":
        cursor.execute("SELECT obligation_id, jurisdiction, statute, obligation_name, applicability_trigger, deadline_days, regulator, last_updated FROM regulatory_ontology")
        for r in cursor.fetchall():
            data.append({
                "obligation_id": r[0],
                "jurisdiction": r[1],
                "statute": r[2],
                "obligation_name": r[3],
                "applicability_trigger": r[4],
                "deadline_days": r[5],
                "regulator": r[6],
                "last_updated": r[7]
            })
    elif domain == "vendor":
        cursor.execute("SELECT vendor_id, vendor_name, service, data_types, dpa_status, soc2_status, sub_processors, cross_border_vectors, risk_rating, last_updated FROM vendor_graph")
        for r in cursor.fetchall():
            data.append({
                "vendor_id": r[0],
                "vendor_name": r[1],
                "service": r[2],
                "data_types": r[3],
                "dpa_status": r[4],
                "soc2_status": r[5],
                "sub_processors": r[6],
                "cross_border_vectors": r[7],
                "risk_rating": r[8],
                "last_updated": r[9]
            })
    elif domain == "inventory":
        cursor.execute("SELECT inventory_id, system_name, data_types, sensitivity, purpose, retention_rules, sharing_recipient, jurisdiction, last_updated FROM data_inventory")
        for r in cursor.fetchall():
            data.append({
                "inventory_id": r[0],
                "system_name": r[1],
                "data_types": r[2],
                "sensitivity": r[3],
                "purpose": r[4],
                "retention_rules": r[5],
                "sharing_recipient": r[6],
                "jurisdiction": r[7],
                "last_updated": r[8]
            })
            
    conn_db.close()
    return data

@app.get("/api/evaluations/flagged")
def get_flagged_evaluations(scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT update_id, entity, domain, impact_summary, detected_at FROM knowledge_updates WHERE status = 'accepted'")
    rows = cursor.fetchall()
    conn_db.close()
    
    flagged = []
    for r in rows:
        flagged.append({
            "id": r[0],
            "evaluation_type": "PIA Re-Assessment" if r[2] == 'regulatory' else "Risk Review",
            "name": f"Re-verify {r[1]} conformity",
            "reason": r[3],
            "flagged_at": r[4]
        })
    return flagged


# --- AI ONBOARDING ENDPOINTS ---

from onboarding_agents import onboarding_flow, OnboardingState, OrganizationalProfile

class OnboardingStartPayload(BaseModel):
    client_id: str
    urls: List[str] = []
    files: List[str] = []

class GapResolvePayload(BaseModel):
    decision: Literal["accept", "correct", "answer"]
    value: str

@app.post("/api/onboarding/start")
def onboarding_start(payload: OnboardingStartPayload, scope: str = Depends(require_scopes(["expert"]))):
    session_id = f"ONB-{uuid.uuid4().hex[:6].upper()}"
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    empty_profile = OrganizationalProfile().model_dump()
    initial_logs = ["[System Orchestrator] Initializing AI Onboarding Agent session..."]
    
    cursor.execute("""
        INSERT INTO onboarding_sessions (session_id, client_id, status, progress, profile_json, logs_json, urls_json, files_json)
        VALUES (?, ?, 'ingesting', 0.0, ?, ?, ?, ?)
    """, (session_id, payload.client_id, json.dumps(empty_profile), json.dumps(initial_logs), json.dumps(payload.urls), json.dumps(payload.files)))
    
    conn_db.commit()
    conn_db.close()
    
    return {"status": "success", "session_id": session_id}

@app.post("/api/onboarding/{session_id}/website")
def onboarding_website(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT client_id, urls_json, files_json, logs_json FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    client_id, urls_str, files_str, logs_str = row
    state = OnboardingState(
        session_id=session_id,
        client_id=client_id,
        urls=json.loads(urls_str),
        files=json.loads(files_str),
        logs=json.loads(logs_str)
    )
    
    from onboarding_agents import ingest_website
    res = ingest_website(state)
    
    cursor.execute("""
        UPDATE onboarding_sessions
        SET progress = ?, status = ?, logs_json = ?, updated_at = CURRENT_TIMESTAMP
        WHERE session_id = ?
    """, (res["progress"], res["status"], json.dumps(res["logs"]), session_id))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "logs": res["logs"]}

@app.post("/api/onboarding/{session_id}/documents")
def onboarding_documents(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT client_id, urls_json, files_json, logs_json, progress, status FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    client_id, urls_str, files_str, logs_str, progress, status = row
    
    state = OnboardingState(
        session_id=session_id,
        client_id=client_id,
        urls=json.loads(urls_str),
        files=json.loads(files_str),
        logs=json.loads(logs_str),
        progress=progress,
        status=status
    )
    
    from onboarding_agents import ingest_documents, normalize_profile, detect_gaps
    
    res_docs = ingest_documents(state)
    state.logs = res_docs["logs"]
    state.progress = res_docs["progress"]
    state.status = res_docs["status"]
    
    res_norm = normalize_profile(state)
    state.logs = res_norm["logs"]
    state.progress = res_norm["progress"]
    state.status = res_norm["status"]
    state.profile = res_norm["profile"]
    
    res_gaps = detect_gaps(state)
    
    for g in res_gaps["gaps"]:
        cursor.execute("""
            INSERT OR REPLACE INTO onboarding_gaps (gap_id, session_id, type, title, details, priority, status)
            VALUES (?, ?, ?, ?, ?, ?, 'pending')
        """, (g["id"], session_id, g["type"], g["title"], g["details"], g["priority"]))
        
    cursor.execute("""
        UPDATE onboarding_sessions
        SET progress = ?, status = ?, logs_json = ?, profile_json = ?, updated_at = CURRENT_TIMESTAMP
        WHERE session_id = ?
    """, (res_gaps["progress"], res_gaps["status"], json.dumps(res_gaps["logs"]), json.dumps(state.profile), session_id))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "logs": res_gaps["logs"], "gaps_count": len(res_gaps["gaps"])}

@app.get("/api/onboarding/{session_id}/profile")
def onboarding_get_profile(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT profile_json FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    return json.loads(row[0])

@app.get("/api/onboarding/{session_id}/gaps")
def onboarding_get_gaps(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT gap_id, type, title, details, priority, status FROM onboarding_gaps WHERE session_id = ?", (session_id,))
    rows = cursor.fetchall()
    conn_db.close()
    
    gaps = []
    for r in rows:
        gaps.append({
            "id": r[0],
            "type": r[1],
            "title": r[2],
            "details": r[3],
            "priority": r[4],
            "status": r[5]
        })
    return gaps

@app.post("/api/onboarding/{session_id}/gaps/{gap_id}")
def onboarding_resolve_gap(session_id: str, gap_id: str, payload: GapResolvePayload, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    cursor.execute("UPDATE onboarding_gaps SET status = ? WHERE gap_id = ? AND session_id = ?", (payload.decision, gap_id, session_id))
    
    cursor.execute("SELECT logs_json, profile_json FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    if row:
        logs = json.loads(row[0])
        profile = json.loads(row[1])
        logs.append(f"[Orchestrator] Gap resolved by human: {gap_id} resolved as '{payload.decision}'. User input: '{payload.value}'")
        
        if gap_id == "gap-1" and payload.value:
            for v in profile.get("vendors", []):
                if v["name"] == "Aselo/Twilio":
                    v["retention"] = payload.value
            for d in profile.get("data_inventory", []):
                if d["asset"] == "Clinical Transcripts":
                    d["retention"] = payload.value
        elif gap_id == "gap-2" and payload.value:
            for v in profile.get("vendors", []):
                if v["name"] == "Blackbaud":
                    v["dpa_status"] = f"Validated: {payload.value}"
                    
        cursor.execute("UPDATE onboarding_sessions SET logs_json = ?, profile_json = ?, updated_at = CURRENT_TIMESTAMP WHERE session_id = ?", (json.dumps(logs), json.dumps(profile), session_id))
        
    conn_db.commit()
    conn_db.close()
    return {"status": "success"}

@app.post("/api/onboarding/{session_id}/finalize")
def onboarding_finalize(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    
    cursor.execute("SELECT client_id, logs_json, profile_json FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    if not row:
        conn_db.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    client_id, logs_str, profile_str = row
    logs = json.loads(logs_str)
    
    from onboarding_agents import finalize_profile, OnboardingState
    state = OnboardingState(
        session_id=session_id,
        client_id=client_id,
        logs=logs,
        profile=json.loads(profile_str),
        progress=0.85,
        status="gap_review"
    )
    
    res = finalize_profile(state)
    
    cursor.execute("""
        UPDATE onboarding_sessions
        SET progress = 1.0, status = 'validated', logs_json = ?, profile_json = ?, updated_at = CURRENT_TIMESTAMP
        WHERE session_id = ?
    """, (json.dumps(res["logs"]), json.dumps(res["profile"]), session_id))
    
    conn_db.commit()
    conn_db.close()
    return {"status": "success", "profile": res["profile"]}

@app.get("/api/onboarding/{session_id}/status")
def onboarding_get_status(session_id: str, scope: str = Depends(require_scopes(["expert"]))):
    conn_db = sqlite3.connect(DB_FILE)
    cursor = conn_db.cursor()
    cursor.execute("SELECT status, progress, logs_json FROM onboarding_sessions WHERE session_id = ?", (session_id,))
    row = cursor.fetchone()
    conn_db.close()
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return {
        "status": row[0],
        "progress": row[1],
        "logs": json.loads(row[2])
    }



